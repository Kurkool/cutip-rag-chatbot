"""Tests for ingestion v2 (Opus 4.7-first universal pipeline)."""
import pytest

from ingest.services import ingestion_v2


def test_module_imports():
    assert hasattr(ingestion_v2, "ingest_v2")
    assert hasattr(ingestion_v2, "ensure_pdf")
    assert hasattr(ingestion_v2, "extract_hyperlinks")
    assert hasattr(ingestion_v2, "opus_parse_and_chunk")


def test_ensure_pdf_passthrough_for_pdf_input():
    # Minimal valid PDF bytes (header only — PyMuPDF can still open this for a sanity round-trip).
    pdf_bytes = b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n1 0 obj\n<< /Type /Catalog >>\nendobj\n%%EOF"
    result = ingestion_v2.ensure_pdf(pdf_bytes, "doc.pdf")
    assert result is pdf_bytes  # identity check: no copy, no re-encode


def test_ensure_pdf_invokes_libreoffice_for_docx(monkeypatch):
    """DOCX input routes through _convert_to_pdf (no real LibreOffice call)."""
    calls = {}

    def fake_convert(blob, ext):
        calls["blob"] = blob
        calls["ext"] = ext
        return b"%PDF-1.4\nconverted\n%%EOF"

    import ingest.services.ingest_helpers as v1_mod
    monkeypatch.setattr(v1_mod, "_convert_to_pdf", fake_convert)

    out = ingestion_v2.ensure_pdf(b"DOCX-fake-bytes", "form.docx")

    assert out == b"%PDF-1.4\nconverted\n%%EOF"
    assert calls["blob"] == b"DOCX-fake-bytes"
    assert calls["ext"] == ".docx"


def test_ensure_pdf_rejects_unsupported_extension():
    with pytest.raises(ValueError, match="unsupported extension"):
        ingestion_v2.ensure_pdf(b"junk", "image.jpg")


def test_extract_hyperlinks_returns_per_page_uris():
    """Synthetic PDF with one hyperlink annotation → one sidecar entry."""
    import pymupdf

    doc = pymupdf.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Click here for details")
    rect = pymupdf.Rect(72, 70, 200, 90)
    page.insert_link({"kind": pymupdf.LINK_URI, "from": rect, "uri": "https://example.org/details"})
    pdf_bytes = doc.tobytes()
    doc.close()

    links = ingestion_v2.extract_hyperlinks(pdf_bytes)

    assert len(links) == 1
    assert links[0]["page"] == 1
    assert links[0]["uri"] == "https://example.org/details"
    assert "Click here" in links[0]["text"]


def test_extract_hyperlinks_empty_pdf_returns_empty_list():
    import pymupdf

    doc = pymupdf.open()
    doc.new_page()
    pdf_bytes = doc.tobytes()
    doc.close()

    links = ingestion_v2.extract_hyperlinks(pdf_bytes)

    assert links == []


def test_extract_hyperlinks_skips_uris_already_visible_in_text():
    """URIs that already appear as visible text should not be duplicated in the sidecar.

    v2's Opus prompt tells the model that sidecar URIs are 'hidden in PDF
    annotations and NOT visible on the rendered page'. If we leak a visible
    URI into the sidecar, we break that contract and Opus emits duplicate
    markdown links.
    """
    import pymupdf

    doc = pymupdf.open()
    page = doc.new_page()
    # The URL is written as visible text on the page.
    page.insert_text((72, 72), "Visit https://example.org/visible for details")
    # And there's also a link annotation pointing to the SAME URL.
    rect = pymupdf.Rect(72, 70, 400, 90)
    page.insert_link({
        "kind": pymupdf.LINK_URI,
        "from": rect,
        "uri": "https://example.org/visible",
    })
    # Plus a SECOND link whose URI is NOT visible in the text — this one must survive.
    rect2 = pymupdf.Rect(72, 100, 400, 120)
    page.insert_text((72, 102), "Click here for more")
    page.insert_link({
        "kind": pymupdf.LINK_URI,
        "from": rect2,
        "uri": "https://example.org/hidden",
    })
    pdf_bytes = doc.tobytes()
    doc.close()

    links = ingestion_v2.extract_hyperlinks(pdf_bytes)

    # Only the hidden one should be in the sidecar.
    assert len(links) == 1
    assert links[0]["uri"] == "https://example.org/hidden"


@pytest.mark.asyncio
async def test_opus_parse_and_chunk_happy_path(monkeypatch):
    """Mocked Opus returns a valid tool call → list of Document chunks."""
    from langchain_core.messages import AIMessage

    async def fake_ainvoke(messages):
        return AIMessage(
            content="",
            tool_calls=[{
                "name": "record_chunks",
                "args": {
                    "chunks": [
                        {"text": "chunk one body", "section_path": "Intro", "page": 1, "has_table": False},
                        {"text": "chunk two body", "section_path": "", "page": 2, "has_table": True},
                    ],
                },
                "id": "call_1",
            }],
        )

    class FakeLLM:
        def bind_tools(self, tools, tool_choice=None):
            return self

        async def ainvoke(self, messages):
            return await fake_ainvoke(messages)

    # Patch the LLM factory used by v2.
    monkeypatch.setattr(ingestion_v2, "_get_opus_llm", lambda: FakeLLM())

    pdf_bytes = b"%PDF-1.4\nminimal\n%%EOF"
    hyperlinks = [{"page": 1, "text": "ref", "uri": "https://example.org"}]

    chunks = await ingestion_v2.opus_parse_and_chunk(pdf_bytes, hyperlinks, "test.pdf")

    assert len(chunks) == 2
    assert chunks[0].page_content == "chunk one body"
    assert chunks[0].metadata["section_path"] == "Intro"
    assert chunks[0].metadata["page"] == 1
    assert chunks[0].metadata["has_table"] is False
    assert chunks[1].metadata["has_table"] is True


@pytest.mark.asyncio
async def test_opus_parse_and_chunk_filters_refusal_chunks(monkeypatch):
    """Chunks whose text matches a known refusal pattern are dropped."""
    from langchain_core.messages import AIMessage

    class FakeLLM:
        def bind_tools(self, tools, tool_choice=None):
            return self

        async def ainvoke(self, messages):
            return AIMessage(content="", tool_calls=[{
                "name": "record_chunks",
                "args": {
                    "chunks": [
                        {"text": "I cannot process this document — please re-upload.", "page": 1},
                        {"text": "real substantive content about วิทยานิพนธ์", "page": 1},
                    ],
                },
                "id": "c1",
            }])

    monkeypatch.setattr(ingestion_v2, "_get_opus_llm", lambda: FakeLLM())

    chunks = await ingestion_v2.opus_parse_and_chunk(b"%PDF-1.4\n%%EOF", [], "t.pdf")

    assert len(chunks) == 1
    assert "real substantive content" in chunks[0].page_content


@pytest.mark.asyncio
async def test_opus_parse_and_chunk_returns_empty_when_no_tool_call(monkeypatch):
    """If Opus replies without calling the tool, we return [] and do not crash."""
    from langchain_core.messages import AIMessage

    class FakeLLM:
        def bind_tools(self, tools, tool_choice=None):
            return self

        async def ainvoke(self, messages):
            # Note: no tool_calls attribute populated (model did not call the tool)
            return AIMessage(content="I have nothing to say.")

    monkeypatch.setattr(ingestion_v2, "_get_opus_llm", lambda: FakeLLM())

    chunks = await ingestion_v2.opus_parse_and_chunk(b"%PDF-1.4\n%%EOF", [], "t.pdf")

    assert chunks == []


@pytest.mark.asyncio
async def test_ingest_v2_orchestrates_pipeline(monkeypatch):
    """End-to-end: ensure_pdf → extract_hyperlinks → opus → _upsert."""
    from langchain_core.documents import Document

    calls = {}

    def fake_ensure_pdf(blob, fn):
        calls["ensure_pdf"] = (blob, fn)
        return b"%PDF-1.4\nnormalized\n%%EOF"

    def fake_extract_page_text(pdf):
        # Return non-empty text so the OCR path is NOT triggered.
        return {1: "some text content"}

    def fake_extract_hyperlinks(pdf):
        calls["extract_hyperlinks"] = pdf
        return [{"page": 1, "text": "t", "uri": "https://x"}]

    async def fake_opus(pdf, links, fn, ocr_sidecar=None):
        calls["opus"] = (pdf, links, fn)
        return [Document(page_content="body", metadata={"page": 1, "section_path": "", "has_table": False})]

    async def fake_upsert(chunks, namespace, extra_metadata):
        calls["upsert"] = {
            "chunks_len": len(chunks),
            "namespace": namespace,
            "extra_metadata": extra_metadata,
        }
        return len(chunks)

    monkeypatch.setattr(ingestion_v2, "ensure_pdf", fake_ensure_pdf)
    monkeypatch.setattr(ingestion_v2, "extract_page_text", fake_extract_page_text)
    monkeypatch.setattr(ingestion_v2, "extract_hyperlinks", fake_extract_hyperlinks)
    monkeypatch.setattr(ingestion_v2, "opus_parse_and_chunk", fake_opus)

    import ingest.services.ingest_helpers as v1_mod
    monkeypatch.setattr(v1_mod, "_upsert", fake_upsert)

    result = await ingestion_v2.ingest_v2(
        file_bytes=b"DOCX-bytes",
        filename="form.docx",
        namespace="cutip_v2_audit",
        tenant_id="cutip_01",
        doc_category="form",
        download_link="https://drive.google.com/file/d/xyz/view",
    )

    assert result == 1
    assert calls["ensure_pdf"] == (b"DOCX-bytes", "form.docx")
    assert calls["extract_hyperlinks"] == b"%PDF-1.4\nnormalized\n%%EOF"
    assert calls["opus"][2] == "form.docx"
    assert calls["upsert"]["namespace"] == "cutip_v2_audit"
    meta = calls["upsert"]["extra_metadata"]
    assert meta["tenant_id"] == "cutip_01"
    assert meta["source_filename"] == "form.docx"
    assert meta["doc_category"] == "form"
    assert meta["source_type"] == "pdf"  # everything normalizes to pdf in v2
    assert meta["download_link"] == "https://drive.google.com/file/d/xyz/view"


def test_extract_page_text_returns_per_page_dict(tiny_text_pdf_bytes):
    result = ingestion_v2.extract_page_text(tiny_text_pdf_bytes)
    assert set(result.keys()) == {1}
    assert "hello" in result[1]


def test_extract_page_text_pure_scan_returns_all_empty(pure_scan_pdf_bytes):
    result = ingestion_v2.extract_page_text(pure_scan_pdf_bytes)
    assert set(result.keys()) == {1, 2}
    assert all(v == "" for v in result.values())
    assert sum(len(v) for v in result.values()) == 0


def test_get_ocr_client_is_cached_async_anthropic():
    from anthropic import AsyncAnthropic

    # Clear the cache from any prior test run so we verify caching in isolation.
    ingestion_v2._get_ocr_client.cache_clear()

    c1 = ingestion_v2._get_ocr_client()
    c2 = ingestion_v2._get_ocr_client()

    assert isinstance(c1, AsyncAnthropic)
    assert c1 is c2  # lru_cache returns the same instance


@pytest.mark.asyncio
async def test_ocr_pdf_pages_returns_per_page_dict(pure_scan_pdf_bytes, monkeypatch):
    """Mock anthropic client returns scripted text per page → dict[int,str]."""
    from unittest.mock import AsyncMock, MagicMock

    class FakeContentBlock:
        def __init__(self, text):
            self.type = "text"
            self.text = text

    call_count = {"n": 0}

    async def fake_create(**kwargs):
        call_count["n"] += 1
        resp = MagicMock()
        resp.content = [FakeContentBlock(f"page {call_count['n']} ocr text")]
        return resp

    fake_client = MagicMock()
    fake_client.messages = MagicMock()
    fake_client.messages.create = AsyncMock(side_effect=fake_create)

    ingestion_v2._get_ocr_client.cache_clear()
    monkeypatch.setattr(ingestion_v2, "_get_ocr_client", lambda: fake_client)

    result = await ingestion_v2.ocr_pdf_pages(pure_scan_pdf_bytes, "test.pdf")

    assert set(result.keys()) == {1, 2}
    assert "ocr text" in result[1]
    assert "ocr text" in result[2]
    assert call_count["n"] == 2  # one call per page


@pytest.mark.asyncio
async def test_ocr_pdf_pages_partial_failure_returns_empty_string(pure_scan_pdf_bytes, monkeypatch):
    from unittest.mock import AsyncMock, MagicMock

    class FakeContentBlock:
        def __init__(self, text):
            self.type = "text"
            self.text = text

    # First call raises; second succeeds. MagicMock's side_effect iterates.
    call_count = {"n": 0}

    async def fake_create(**kwargs):
        call_count["n"] += 1
        if call_count["n"] == 1:
            raise RuntimeError("simulated rate limit")
        resp = MagicMock()
        resp.content = [FakeContentBlock("page two ok")]
        return resp

    fake_client = MagicMock()
    fake_client.messages = MagicMock()
    fake_client.messages.create = AsyncMock(side_effect=fake_create)

    ingestion_v2._get_ocr_client.cache_clear()
    monkeypatch.setattr(ingestion_v2, "_get_ocr_client", lambda: fake_client)

    result = await ingestion_v2.ocr_pdf_pages(pure_scan_pdf_bytes, "test.pdf")

    # One page failed → empty string at that key, no raise.
    assert "" in result.values()
    assert any("page two" in v for v in result.values())


@pytest.mark.asyncio
async def test_ocr_pdf_pages_all_pages_fail_raises(pure_scan_pdf_bytes, monkeypatch):
    from unittest.mock import AsyncMock, MagicMock

    async def fake_create(**kwargs):
        raise RuntimeError("every call fails")

    fake_client = MagicMock()
    fake_client.messages = MagicMock()
    fake_client.messages.create = AsyncMock(side_effect=fake_create)

    ingestion_v2._get_ocr_client.cache_clear()
    monkeypatch.setattr(ingestion_v2, "_get_ocr_client", lambda: fake_client)

    with pytest.raises(RuntimeError, match="OCR failed for all pages"):
        await ingestion_v2.ocr_pdf_pages(pure_scan_pdf_bytes, "test.pdf")


def test_format_ocr_sidecar_empty_dict_returns_placeholder():
    from ingest.services._v2_prompts import format_ocr_sidecar
    assert format_ocr_sidecar({}) == "(no OCR sidecar — document text layer sufficient)"


def test_format_ocr_sidecar_populated_renders_per_page_sections():
    from ingest.services._v2_prompts import format_ocr_sidecar
    out = format_ocr_sidecar({1: "page one text", 2: "page two\nmore text"})
    assert "### Page 1" in out
    assert "page one text" in out
    assert "### Page 2" in out
    assert "page two\nmore text" in out


def test_user_prompt_template_has_ocr_block_placeholder():
    from ingest.services._v2_prompts import USER_PROMPT_TEMPLATE
    assert "{ocr_block}" in USER_PROMPT_TEMPLATE


@pytest.mark.asyncio
async def test_opus_parse_and_chunk_injects_ocr_sidecar(monkeypatch, pure_scan_pdf_bytes):
    """When ocr_sidecar is provided, the user message must contain the formatted OCR block."""
    from unittest.mock import AsyncMock, MagicMock

    captured: dict = {}

    fake_llm = MagicMock()

    async def fake_ainvoke(messages):
        # Capture the human message content for assertion.
        human = messages[1]
        captured["content"] = human.content
        return MagicMock(tool_calls=[{"args": {"chunks": []}}])

    fake_llm.ainvoke = AsyncMock(side_effect=fake_ainvoke)
    fake_llm.bind_tools = MagicMock(return_value=fake_llm)

    ingestion_v2._get_opus_llm.cache_clear()
    monkeypatch.setattr(ingestion_v2, "_get_opus_llm", lambda: fake_llm)

    await ingestion_v2.opus_parse_and_chunk(
        pdf_bytes=pure_scan_pdf_bytes,
        hyperlinks=[],
        filename="x.pdf",
        ocr_sidecar={1: "hello from page 1", 2: "hello from page 2"},
    )

    # Find the text content block in the human message.
    text_blocks = [b["text"] for b in captured["content"] if b.get("type") == "text"]
    assert any("hello from page 1" in t for t in text_blocks)
    assert any("### Page 1" in t for t in text_blocks)


@pytest.mark.asyncio
async def test_opus_parse_and_chunk_without_ocr_sidecar_uses_placeholder(monkeypatch, pure_scan_pdf_bytes):
    from unittest.mock import AsyncMock, MagicMock

    captured: dict = {}
    fake_llm = MagicMock()

    async def fake_ainvoke(messages):
        captured["content"] = messages[1].content
        return MagicMock(tool_calls=[{"args": {"chunks": []}}])

    fake_llm.ainvoke = AsyncMock(side_effect=fake_ainvoke)
    fake_llm.bind_tools = MagicMock(return_value=fake_llm)

    ingestion_v2._get_opus_llm.cache_clear()
    monkeypatch.setattr(ingestion_v2, "_get_opus_llm", lambda: fake_llm)

    await ingestion_v2.opus_parse_and_chunk(
        pdf_bytes=pure_scan_pdf_bytes,
        hyperlinks=[],
        filename="x.pdf",
        ocr_sidecar=None,
    )

    text_blocks = [b["text"] for b in captured["content"] if b.get("type") == "text"]
    assert any("no OCR sidecar" in t for t in text_blocks)


@pytest.mark.asyncio
async def test_ingest_v2_pure_scan_triggers_ocr_path(monkeypatch, pure_scan_pdf_bytes):
    """Pure-scan PDF must trigger ocr_pdf_pages and pass its output downstream."""
    from unittest.mock import AsyncMock

    captured: dict = {}

    async def fake_ocr(pdf_bytes, filename):
        captured["ocr_called"] = True
        captured["filename"] = filename
        return {1: "ocr p1", 2: "ocr p2"}

    async def fake_opus(pdf_bytes, hyperlinks, filename, ocr_sidecar=None):
        captured["ocr_sidecar_arg"] = ocr_sidecar
        return []  # 0 chunks — we're not exercising the upsert here

    monkeypatch.setattr(ingestion_v2, "ocr_pdf_pages", fake_ocr)
    monkeypatch.setattr(ingestion_v2, "opus_parse_and_chunk", fake_opus)

    result = await ingestion_v2.ingest_v2(
        file_bytes=pure_scan_pdf_bytes,
        filename="scan.pdf",
        namespace="ns-test",
        tenant_id="tenant_x",
    )

    assert result == 0
    assert captured.get("ocr_called") is True
    assert captured["ocr_sidecar_arg"] == {1: "ocr p1", 2: "ocr p2"}


@pytest.mark.asyncio
async def test_ingest_v2_text_layer_skips_ocr(monkeypatch, tiny_text_pdf_bytes):
    """Text-layer PDF must NOT trigger ocr_pdf_pages."""
    from unittest.mock import AsyncMock

    captured: dict = {}

    async def fake_ocr(pdf_bytes, filename):
        captured["ocr_called"] = True
        return {}

    async def fake_opus(pdf_bytes, hyperlinks, filename, ocr_sidecar=None):
        captured["ocr_sidecar_arg"] = ocr_sidecar
        return []

    monkeypatch.setattr(ingestion_v2, "ocr_pdf_pages", fake_ocr)
    monkeypatch.setattr(ingestion_v2, "opus_parse_and_chunk", fake_opus)

    await ingestion_v2.ingest_v2(
        file_bytes=tiny_text_pdf_bytes,
        filename="text.pdf",
        namespace="ns-test",
        tenant_id="tenant_x",
    )

    assert captured.get("ocr_called") is not True
    assert captured["ocr_sidecar_arg"] is None


def test_read_ocr_docx_as_pages_splits_on_page_headings():
    """Docx with 'หน้า N' level-2 headings → pages keyed by N with joined paragraph text."""
    import io
    from docx import Document as DocxDocument

    docx = DocxDocument()
    docx.add_heading("doc title", level=1)            # ignored (level 1)
    docx.add_heading("หน้า 1", level=2)
    docx.add_paragraph("first para of page 1")
    docx.add_paragraph("second para of page 1")
    docx.add_heading("หน้า 2", level=2)
    docx.add_paragraph("only para of page 2")
    buf = io.BytesIO()
    docx.save(buf)

    result = ingestion_v2._read_ocr_docx_as_pages(buf.getvalue())

    assert result == {
        1: "first para of page 1\nsecond para of page 1",
        2: "only para of page 2",
    }


def test_read_ocr_docx_as_pages_fallback_single_page_when_no_markers():
    """Docx without any 'หน้า N' heading → all paragraphs join into page 1."""
    import io
    from docx import Document as DocxDocument

    docx = DocxDocument()
    docx.add_paragraph("alpha")
    docx.add_paragraph("beta")
    docx.add_paragraph("gamma")
    buf = io.BytesIO()
    docx.save(buf)

    result = ingestion_v2._read_ocr_docx_as_pages(buf.getvalue())

    assert result == {1: "alpha\nbeta\ngamma"}


def test_read_ocr_docx_as_pages_raises_on_corrupt_bytes():
    """python-docx's PackageNotFoundError is wrapped in ValueError for callers."""
    with pytest.raises(ValueError, match="not a valid .ocr.docx"):
        ingestion_v2._read_ocr_docx_as_pages(b"this is not a zip file")


def test_read_ocr_docx_as_pages_raises_on_zip_without_content_types():
    """A valid ZIP that lacks the OPC [Content_Types].xml structure also wraps to ValueError."""
    import io
    import zipfile

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as zf:
        zf.writestr("some_random_file.txt", "not a docx")
    with pytest.raises(ValueError, match="not a valid .ocr.docx"):
        ingestion_v2._read_ocr_docx_as_pages(buf.getvalue())


def test_format_pages_for_text_only_empty_returns_placeholder():
    result = ingestion_v2._format_pages_for_text_only({})
    assert result == "(empty document)"


def test_format_pages_for_text_only_has_page_markers_and_content():
    result = ingestion_v2._format_pages_for_text_only({
        1: "alpha text",
        2: "beta text",
    })
    # Page markers use "### Page N" (matches format_ocr_sidecar convention).
    assert "### Page 1" in result
    assert "### Page 2" in result
    # Content preserved in order.
    i1 = result.index("alpha text")
    i2 = result.index("beta text")
    assert i1 < i2


def test_user_prompt_template_text_only_has_required_placeholders():
    from ingest.services._v2_prompts import USER_PROMPT_TEMPLATE_TEXT_ONLY
    assert "{filename}" in USER_PROMPT_TEMPLATE_TEXT_ONLY
    assert "{sidecar_block}" in USER_PROMPT_TEMPLATE_TEXT_ONLY
    assert "{page_text_block}" in USER_PROMPT_TEMPLATE_TEXT_ONLY
    # Must NOT tell Opus to parse an attached PDF (there is no PDF block).
    assert "attached PDF" not in USER_PROMPT_TEMPLATE_TEXT_ONLY


@pytest.mark.asyncio
async def test_opus_parse_and_chunk_text_only_omits_document_block(
    monkeypatch, pure_scan_pdf_bytes
):
    """With ocr_sidecar populated, the human message must have ONE text block and NO document block."""
    from unittest.mock import AsyncMock, MagicMock

    captured: dict = {}
    fake_llm = MagicMock()

    async def fake_ainvoke(messages):
        captured["content"] = messages[1].content
        return MagicMock(tool_calls=[{"args": {"chunks": []}}])

    fake_llm.ainvoke = AsyncMock(side_effect=fake_ainvoke)
    fake_llm.bind_tools = MagicMock(return_value=fake_llm)

    ingestion_v2._get_opus_llm.cache_clear()
    monkeypatch.setattr(ingestion_v2, "_get_opus_llm", lambda: fake_llm)

    await ingestion_v2.opus_parse_and_chunk(
        pdf_bytes=pure_scan_pdf_bytes,
        hyperlinks=[],
        filename="x.pdf",
        ocr_sidecar={1: "only page"},
    )

    blocks = captured["content"]
    block_types = [b.get("type") for b in blocks]
    assert block_types == ["text"], f"expected ['text'], got {block_types}"
    assert "only page" in blocks[0]["text"]
    assert "### Page 1" in blocks[0]["text"]


@pytest.mark.asyncio
async def test_opus_parse_and_chunk_text_only_accepts_none_pdf_bytes(
    monkeypatch,
):
    """For .ocr.docx (no PDF at all), pdf_bytes=None must not crash the text-only path."""
    from unittest.mock import AsyncMock, MagicMock

    captured: dict = {}
    fake_llm = MagicMock()

    async def fake_ainvoke(messages):
        captured["content"] = messages[1].content
        return MagicMock(tool_calls=[{"args": {"chunks": []}}])

    fake_llm.ainvoke = AsyncMock(side_effect=fake_ainvoke)
    fake_llm.bind_tools = MagicMock(return_value=fake_llm)

    ingestion_v2._get_opus_llm.cache_clear()
    monkeypatch.setattr(ingestion_v2, "_get_opus_llm", lambda: fake_llm)

    chunks = await ingestion_v2.opus_parse_and_chunk(
        pdf_bytes=None,
        hyperlinks=[],
        filename="y.ocr.docx",
        ocr_sidecar={1: "aaa", 2: "bbb"},
    )

    # Confirms no exception, content is text-only.
    assert chunks == []
    block_types = [b.get("type") for b in captured["content"]]
    assert block_types == ["text"]
    assert "aaa" in captured["content"][0]["text"]
    assert "bbb" in captured["content"][0]["text"]


@pytest.mark.asyncio
async def test_opus_parse_and_chunk_pdf_path_keeps_document_block(
    monkeypatch, pure_scan_pdf_bytes
):
    """Regression guard: text-layer PDF path still sends document + text blocks."""
    from unittest.mock import AsyncMock, MagicMock

    captured: dict = {}
    fake_llm = MagicMock()

    async def fake_ainvoke(messages):
        captured["content"] = messages[1].content
        return MagicMock(tool_calls=[{"args": {"chunks": []}}])

    fake_llm.ainvoke = AsyncMock(side_effect=fake_ainvoke)
    fake_llm.bind_tools = MagicMock(return_value=fake_llm)

    ingestion_v2._get_opus_llm.cache_clear()
    monkeypatch.setattr(ingestion_v2, "_get_opus_llm", lambda: fake_llm)

    await ingestion_v2.opus_parse_and_chunk(
        pdf_bytes=pure_scan_pdf_bytes,
        hyperlinks=[],
        filename="x.pdf",
        ocr_sidecar=None,
    )

    block_types = [b.get("type") for b in captured["content"]]
    assert block_types == ["document", "text"], f"expected document+text, got {block_types}"


@pytest.mark.asyncio
async def test_ingest_v2_ocr_docx_skips_libreoffice_and_hyperlinks(
    monkeypatch,
):
    """Filename ending in .ocr.docx → read docx paragraphs, skip ensure_pdf + extract_hyperlinks."""
    import io
    from docx import Document as DocxDocument
    from unittest.mock import AsyncMock

    # Build a minimal .ocr.docx with 2 pages.
    docx = DocxDocument()
    docx.add_heading("title", level=1)
    docx.add_heading("หน้า 1", level=2)
    docx.add_paragraph("page one body")
    docx.add_heading("หน้า 2", level=2)
    docx.add_paragraph("page two body")
    buf = io.BytesIO()
    docx.save(buf)
    docx_bytes = buf.getvalue()

    captured: dict = {"ensure_pdf_called": False, "hyperlinks_called": False}

    def fake_ensure_pdf(blob, fn):
        captured["ensure_pdf_called"] = True
        return b""

    def fake_extract_hyperlinks(pdf):
        captured["hyperlinks_called"] = True
        return []

    async def fake_opus(pdf_bytes, hyperlinks, filename, ocr_sidecar=None):
        captured["opus_pdf_bytes"] = pdf_bytes
        captured["opus_hyperlinks"] = hyperlinks
        captured["opus_ocr_sidecar"] = ocr_sidecar
        return []

    async def fake_upsert(chunks, namespace, extra_metadata):
        return len(chunks)

    monkeypatch.setattr(ingestion_v2, "ensure_pdf", fake_ensure_pdf)
    monkeypatch.setattr(ingestion_v2, "extract_hyperlinks", fake_extract_hyperlinks)
    monkeypatch.setattr(ingestion_v2, "opus_parse_and_chunk", fake_opus)

    import ingest.services.ingest_helpers as v1_mod
    monkeypatch.setattr(v1_mod, "_upsert", fake_upsert)

    result = await ingestion_v2.ingest_v2(
        file_bytes=docx_bytes,
        filename="ประกาศ.ocr.docx",
        namespace="ns-test",
        tenant_id="t",
    )

    assert result == 0
    assert captured["ensure_pdf_called"] is False
    assert captured["hyperlinks_called"] is False
    assert captured["opus_pdf_bytes"] is None
    assert captured["opus_hyperlinks"] == []
    assert captured["opus_ocr_sidecar"] == {
        1: "page one body",
        2: "page two body",
    }


def test_extract_json_from_fence_parses_basic_fence():
    text = """some chatter
```json
{"chunks": [{"text": "a", "page": 1}]}
```
trailing prose"""
    result = ingestion_v2._extract_json_from_fence(text)
    assert result == {"chunks": [{"text": "a", "page": 1}]}


def test_extract_json_from_fence_parses_bare_json_with_prose():
    text = 'Here is the result: {"chunks": [{"text": "alpha", "page": 1}]} done.'
    result = ingestion_v2._extract_json_from_fence(text)
    assert result == {"chunks": [{"text": "alpha", "page": 1}]}


def test_extract_json_from_fence_returns_none_on_malformed():
    text = "```json\nthis is not json at all\n```"
    assert ingestion_v2._extract_json_from_fence(text) is None


def test_extract_json_from_fence_returns_none_on_empty():
    assert ingestion_v2._extract_json_from_fence("") is None
    assert ingestion_v2._extract_json_from_fence(None) is None


def test_text_from_response_handles_string_content():
    from unittest.mock import MagicMock
    fake = MagicMock()
    fake.content = "hello world"
    assert ingestion_v2._text_from_response(fake) == "hello world"


def test_text_from_response_concatenates_text_blocks_only():
    from unittest.mock import MagicMock
    fake = MagicMock()
    fake.content = [
        {"type": "thinking", "thinking": "internal monologue"},  # ignored
        {"type": "text", "text": "first"},
        {"type": "text", "text": "second"},
    ]
    assert ingestion_v2._text_from_response(fake) == "first\nsecond"


def test_pdf_to_image_blocks_returns_image_block_per_page():
    """3-page PDF → 3 image blocks, each base64 PNG with media_type=image/png."""
    import pymupdf

    doc = pymupdf.open()
    for _ in range(3):
        page = doc.new_page()
        page.insert_text((72, 72), "hello")
    pdf_bytes = doc.tobytes()
    doc.close()

    blocks = ingestion_v2._pdf_to_image_blocks(pdf_bytes)

    assert len(blocks) == 3
    for b in blocks:
        assert b["type"] == "image"
        assert b["source"]["type"] == "base64"
        assert b["source"]["media_type"] == "image/png"
        assert isinstance(b["source"]["data"], str)
        assert len(b["source"]["data"]) > 0


def test_pdf_to_image_blocks_raises_on_over_100_pages():
    """Anthropic limit is 100 image blocks per request; surface a clear error."""
    import pymupdf
    import pytest

    doc = pymupdf.open()
    for _ in range(101):
        doc.new_page()
    pdf_bytes = doc.tobytes()
    doc.close()

    with pytest.raises(ValueError, match="exceeds Anthropic's 100-image limit"):
        ingestion_v2._pdf_to_image_blocks(pdf_bytes)
