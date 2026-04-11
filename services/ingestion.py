"""God-mode document ingestion: Vision-first, handles every edge case."""

import asyncio
import io
import logging
import os
import re
import subprocess
import tempfile
from typing import Any

import pandas as pd
import pymupdf
from docx import Document as DocxDocument
from langchain_anthropic import ChatAnthropic
from langchain_core.documents import Document
from langchain_text_splitters import (
    MarkdownHeaderTextSplitter,
    RecursiveCharacterTextSplitter,
)

from config import settings
from services.vectorstore import get_raw_index, get_vectorstore
from services.vision import interpret_spreadsheet, parse_page_image

logger = logging.getLogger(__name__)

# ──────────────────────────────────────
# Smart Chunking Pipeline
# ──────────────────────────────────────

md_header_splitter = MarkdownHeaderTextSplitter(
    headers_to_split_on=[
        ("#", "section"),
        ("##", "subsection"),
        ("###", "topic"),
    ],
    strip_headers=False,
)

text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=settings.CHUNK_SIZE,
    chunk_overlap=settings.CHUNK_OVERLAP,
    separators=["\n\n", "\n", "。", ".", " ", ""],
)


def _smart_chunk(text: str, source: str = "") -> list[Document]:
    """MarkdownHeader split → Recursive split → context enrichment."""
    header_chunks = md_header_splitter.split_text(text)
    if not header_chunks:
        return text_splitter.create_documents(
            [text], metadatas=[{"source": source}]
        )

    final_chunks = text_splitter.split_documents(header_chunks)
    for chunk in final_chunks:
        header_path = " > ".join(
            chunk.metadata[key]
            for key in ["section", "subsection", "topic"]
            if chunk.metadata.get(key)
        )
        if header_path and not chunk.page_content.startswith(header_path):
            chunk.page_content = f"[{header_path}]\n{chunk.page_content}"
        chunk.metadata["source"] = source
    return final_chunks


# ──────────────────────────────────────
# Metadata
# ──────────────────────────────────────

def _build_metadata(
    tenant_id: str,
    source_type: str,
    source_filename: str = "",
    doc_category: str = "general",
    url: str = "",
    download_link: str = "",
) -> dict[str, Any]:
    return {
        "tenant_id": tenant_id,
        "source_type": source_type,
        "source_filename": source_filename,
        "doc_category": doc_category,
        "url": url,
        "download_link": download_link,
    }


# ──────────────────────────────────────
# Legacy format conversion (.doc, .xls, .ppt → PDF)
# ──────────────────────────────────────

def _convert_to_pdf(file_bytes: bytes, src_ext: str) -> bytes:
    """Convert legacy formats to PDF using LibreOffice headless."""
    with tempfile.TemporaryDirectory() as tmpdir:
        src_path = os.path.join(tmpdir, f"input{src_ext}")
        with open(src_path, "wb") as f:
            f.write(file_bytes)

        result = subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", tmpdir, src_path],
            capture_output=True,
            timeout=settings.LIBREOFFICE_TIMEOUT,
        )
        if result.returncode != 0:
            logger.error("LibreOffice failed: %s", result.stderr.decode(errors="replace"))

        pdf_path = os.path.join(tmpdir, "input.pdf")
        if not os.path.exists(pdf_path):
            raise RuntimeError(f"LibreOffice conversion failed for {src_ext}")

        with open(pdf_path, "rb") as f:
            return f.read()


async def ingest_legacy(
    file_bytes: bytes,
    filename: str,
    namespace: str,
    tenant_id: str,
    skip_enrichment: bool = False,
    doc_category: str = "general",
    url: str = "",
    download_link: str = "",
) -> int:
    """Convert .doc/.xls/.ppt → PDF → Vision pipeline."""
    ext = os.path.splitext(filename)[1].lower()
    pdf_bytes = _convert_to_pdf(file_bytes, ext)
    return await ingest_pdf(
        file_bytes=pdf_bytes,
        filename=filename,
        namespace=namespace,
        tenant_id=tenant_id,
        skip_enrichment=skip_enrichment,
        doc_category=doc_category,
        url=url,
        download_link=download_link,
    )


# ──────────────────────────────────────
# Duplicate Detection: delete old vectors before re-ingest
# ──────────────────────────────────────

def _delete_existing_vectors(namespace: str, source_filename: str):
    """Delete all vectors with matching source_filename in the namespace."""
    try:
        index = get_raw_index()
        # Pinecone: delete by metadata filter
        index.delete(
            filter={"source_filename": source_filename},
            namespace=namespace,
        )
        logger.info("Deleted old vectors for '%s' in namespace '%s'", source_filename, namespace)
    except Exception:
        logger.warning("Could not delete old vectors for '%s' (may not exist)", source_filename)


# ──────────────────────────────────────
# PDF Ingestion (Hybrid: text extraction + Vision fallback)
# ──────────────────────────────────────

_VISION_THRESHOLD = settings.PDF_VISION_THRESHOLD
_PDF_BATCH_SIZE = settings.PDF_BATCH_SIZE


async def ingest_pdf(
    file_bytes: bytes,
    filename: str,
    namespace: str,
    tenant_id: str,
    skip_enrichment: bool = False,
    doc_category: str = "general",
    url: str = "",
    download_link: str = "",
) -> int:
    """
    Hybrid PDF ingestion:
    - Pages with enough text → PyMuPDF text extraction (fast, free)
    - Pages with little/no text (tables, forms, scanned) → Claude Vision (accurate)
    - Hidden hyperlinks extracted from all pages

    Handles: text, forms, slides, scanned, tables, diagrams, hidden hyperlinks.
    """
    _delete_existing_vectors(namespace, filename)

    tmp_path = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name

        doc = pymupdf.open(tmp_path)
        if doc.is_encrypted:
            logger.error("Password-protected PDF: %s", filename)
            return 0

        # Phase 1: Extract text + classify each page
        pages_data = []
        vision_pages = []
        for i, page in enumerate(doc):
            text = page.get_text("text").strip()
            hidden_links = _extract_hyperlinks(page)
            page_num = i + 1

            # Has tables? (PyMuPDF can detect table structures)
            has_tables = bool(page.find_tables().tables)

            if len(text) >= _VISION_THRESHOLD and not has_tables:
                # Enough text, no tables → use text extraction (fast)
                if hidden_links:
                    text += "\n\n**Links in this page:**\n" + "\n".join(
                        f"- [{t}]({u})" for t, u in hidden_links
                    )
                pages_data.append({"text": text, "page": page_num})
            else:
                # Low text or has tables → need Vision
                pix = page.get_pixmap(dpi=150)
                vision_pages.append({
                    "img_bytes": pix.tobytes("png"),
                    "links": hidden_links,
                    "page_num": page_num,
                })
        doc.close()

        # Phase 2: Process Vision pages in batches
        logger.info(
            "%s: %d text pages, %d vision pages",
            filename, len(pages_data), len(vision_pages),
        )
        for batch_start in range(0, len(vision_pages), _PDF_BATCH_SIZE):
            if batch_start > 0:
                await asyncio.sleep(2)  # Rate limit: pause between batches
            batch = vision_pages[batch_start:batch_start + _PDF_BATCH_SIZE]
            tasks = [_process_pdf_page(p) for p in batch]
            results = await asyncio.gather(*tasks, return_exceptions=True)
            for result in results:
                if isinstance(result, dict) and result.get("text"):
                    pages_data.append(result)

        # Sort by page number
        pages_data.sort(key=lambda p: p["page"])
        page_texts = pages_data

        if not page_texts:
            return 0

        full_text = "\n\n---\n\n".join(p["text"] for p in page_texts)

        # Detect slides vs document
        avg_chars = len(full_text) / max(len(page_texts), 1)
        is_slides = len(page_texts) > 5 and avg_chars < 500

        if is_slides:
            chunks = _chunk_pages(page_texts, source=filename)
        else:
            chunks = _smart_chunk(full_text, source=filename)
            for chunk in chunks:
                if "page" not in chunk.metadata:
                    chunk.metadata["page"] = 1

        metadata = _build_metadata(
            tenant_id, "pdf", filename, doc_category, url, download_link
        )
        return await _upsert(chunks, namespace, metadata, full_text=full_text, skip_enrichment=skip_enrichment)
    finally:
        if tmp_path and os.path.exists(tmp_path):
            os.unlink(tmp_path)


async def _process_pdf_page(page_data: dict) -> dict:
    """Process a single PDF page: Vision + hyperlinks."""
    try:
        markdown = await parse_page_image(page_data["img_bytes"])
        if page_data["links"]:
            markdown += "\n\n**Links in this page:**\n" + "\n".join(
                f"- [{text}]({uri})" for text, uri in page_data["links"]
            )
        return {"text": markdown, "page": page_data["page_num"]}
    except Exception:
        logger.warning("Failed to process page %d", page_data["page_num"])
        return {"text": "", "page": page_data["page_num"]}


def _extract_hyperlinks(page) -> list[tuple[str, str]]:
    """Extract hidden hyperlinks from a PDF page (not visible as text)."""
    page_text = page.get_text("text")
    links = []
    for link in page.get_links():
        uri = link.get("uri", "")
        if not uri or uri in page_text:
            continue  # Skip if URI is already visible in text
        rect = link.get("from", pymupdf.Rect())
        text = page.get_text("text", clip=rect).strip() or uri
        links.append((text, uri))
    return links


def _chunk_pages(pages: list[dict], source: str = "") -> list[Document]:
    """Page-level chunking for slides: merge short pages."""
    chunks: list[Document] = []
    buffer = ""
    buffer_pages: list[int] = []

    for page in pages:
        text = page["text"].strip()
        if not text:
            continue
        if len(text) < 100:
            buffer += f"\n{text}" if buffer else text
            buffer_pages.append(page["page"])
        else:
            if buffer:
                chunks.append(Document(
                    page_content=buffer,
                    metadata={"source": source, "pages": buffer_pages},
                ))
                buffer = ""
                buffer_pages = []
            page_chunks = _smart_chunk(text, source)
            for c in page_chunks:
                c.metadata["page"] = page["page"]
            chunks.extend(page_chunks)

    if buffer:
        chunks.append(Document(
            page_content=buffer,
            metadata={"source": source, "pages": buffer_pages},
        ))
    return chunks


# ──────────────────────────────────────
# DOCX Ingestion (paragraphs + tables + images)
# ──────────────────────────────────────

async def ingest_docx(
    file_bytes: bytes,
    filename: str,
    namespace: str,
    tenant_id: str,
    skip_enrichment: bool = False,
    doc_category: str = "general",
    url: str = "",
    download_link: str = "",
) -> int:
    """
    DOCX → extract paragraphs + tables + images → combined markdown → chunk → embed.
    No more missing tables or images.
    """
    _delete_existing_vectors(namespace, filename)

    doc = DocxDocument(io.BytesIO(file_bytes))

    parts = []

    # 1. Extract paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            # Detect heading styles
            if para.style and para.style.name.startswith("Heading"):
                level = para.style.name.replace("Heading ", "")
                try:
                    hashes = "#" * int(level)
                except ValueError:
                    hashes = "##"
                parts.append(f"{hashes} {text}")
            else:
                parts.append(text)

    # 2. Extract tables → markdown tables
    for table in doc.tables:
        md_table = _docx_table_to_markdown(table)
        if md_table:
            parts.append(md_table)

    # 3. Extract images → Claude Vision
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            try:
                img_bytes = rel.target_part.blob
                caption = await parse_page_image(img_bytes)
                if caption:
                    parts.append(f"[Image content: {caption}]")
            except Exception:
                logger.warning("Failed to extract image from DOCX '%s'", filename)

    full_text = "\n\n".join(parts)
    chunks = _smart_chunk(full_text, source=filename)
    metadata = _build_metadata(
        tenant_id, "docx", filename, doc_category, url, download_link
    )
    return await _upsert(chunks, namespace, metadata, full_text=full_text, skip_enrichment=skip_enrichment)


def _docx_table_to_markdown(table) -> str:
    """Convert a python-docx table to markdown table."""
    rows = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        rows.append(cells)

    if not rows:
        return ""

    # First row as header
    header = "| " + " | ".join(rows[0]) + " |"
    separator = "| " + " | ".join("---" for _ in rows[0]) + " |"
    body = "\n".join(
        "| " + " | ".join(row) + " |" for row in rows[1:]
    )
    return f"{header}\n{separator}\n{body}"


# ──────────────────────────────────────
# Markdown Ingestion (web content)
# ──────────────────────────────────────

async def ingest_markdown(
    content: str,
    title: str,
    namespace: str,
    tenant_id: str,
    skip_enrichment: bool = False,
    doc_category: str = "general",
    url: str = "",
    download_link: str = "",
) -> int:
    """Markdown content (from Jina Reader) → smart chunk → embed."""
    source_name = title or url or "web_content"
    _delete_existing_vectors(namespace, source_name)

    chunks = _smart_chunk(content, source=source_name)
    metadata = _build_metadata(
        tenant_id, "web", source_name, doc_category, url, download_link
    )
    return await _upsert(chunks, namespace, metadata, full_text=content, skip_enrichment=skip_enrichment)


# ──────────────────────────────────────
# Spreadsheet Ingestion (Claude interprets, batched for large sheets)
# ──────────────────────────────────────

_XLSX_BATCH_ROWS = settings.XLSX_BATCH_ROWS


async def ingest_spreadsheet(
    file_bytes: bytes,
    filename: str,
    namespace: str,
    tenant_id: str,
    skip_enrichment: bool = False,
    doc_category: str = "general",
    url: str = "",
    download_link: str = "",
) -> tuple[int, int]:
    """
    XLSX/CSV → raw dump → Claude interprets (batched for large sheets)
    → markdown → chunk → embed.
    """
    _delete_existing_vectors(namespace, filename)

    is_csv = filename.lower().endswith(".csv")
    total_chunks = 0

    if is_csv:
        df = pd.read_csv(io.BytesIO(file_bytes), header=None)
        structured = await _interpret_dataframe(df)
        chunks = _smart_chunk(structured, source=filename)
        metadata = _build_metadata(
            tenant_id, "spreadsheet", filename, doc_category, url, download_link
        )
        total_chunks = await _upsert(chunks, namespace, metadata, full_text=structured, skip_enrichment=skip_enrichment)
        return 1, total_chunks

    xls = pd.ExcelFile(io.BytesIO(file_bytes))
    sheets_processed = 0
    for sheet_name in xls.sheet_names:
        df = xls.parse(sheet_name, header=None)
        if df.dropna(how="all").empty:
            continue
        structured = await _interpret_dataframe(df, sheet_name=sheet_name)
        chunks = _smart_chunk(structured, source=f"{filename} - {sheet_name}")
        metadata = _build_metadata(
            tenant_id, "spreadsheet", filename, doc_category, url, download_link
        )
        total_chunks += await _upsert(chunks, namespace, metadata, full_text=structured, skip_enrichment=skip_enrichment)
        sheets_processed += 1

    return sheets_processed, total_chunks


async def _interpret_dataframe(df: pd.DataFrame, sheet_name: str = "") -> str:
    """Send DataFrame to Claude in batches for large sheets."""
    df_clean = df.dropna(how="all").dropna(axis=1, how="all")
    total_rows = len(df_clean)

    if total_rows <= _XLSX_BATCH_ROWS:
        # Small sheet: send all at once
        raw = _raw_dataframe_dump(df_clean)
        prefix = f"Sheet: {sheet_name}\n\n" if sheet_name else ""
        return await interpret_spreadsheet(f"{prefix}{raw}")

    # Large sheet: process in batches
    parts = []
    for start in range(0, total_rows, _XLSX_BATCH_ROWS):
        batch = df_clean.iloc[start:start + _XLSX_BATCH_ROWS]
        raw = _raw_dataframe_dump(batch)
        prefix = f"Sheet: {sheet_name} (rows {start+1}-{start+len(batch)})\n\n"
        interpreted = await interpret_spreadsheet(f"{prefix}{raw}")
        parts.append(interpreted)

    return "\n\n".join(parts)


def _raw_dataframe_dump(df: pd.DataFrame) -> str:
    """Dump raw cell data as readable text for Claude to interpret."""
    lines = []
    for i, row in df.iterrows():
        values = []
        for j, val in enumerate(row):
            if pd.notna(val):
                values.append(f"[{j}]={val}")
        if values:
            lines.append(f"Row {i}: {' | '.join(values)}")
    return "\n".join(lines)


# ──────────────────────────────────────
# Contextual Retrieval (Anthropic Research)
# ──────────────────────────────────────

_CONTEXT_PROMPT = (
    "Here is the full document:\n<document>\n{document}\n</document>\n\n"
    "Here is a specific chunk from that document:\n<chunk>\n{chunk}\n</chunk>\n\n"
    "Write a short 1-2 sentence context in the SAME LANGUAGE as the document "
    "that explains where this chunk fits within the document. "
    "Include the document topic, section name, and what this chunk is about. "
    "Reply with ONLY the context, nothing else."
)


async def _enrich_with_context(
    chunks: list[Document], full_text: str
) -> list[Document]:
    """Prepend document-level context to each chunk (~49% retrieval improvement)."""
    if not chunks:
        return chunks

    # Use beginning + end of doc for better context coverage on long documents
    if len(full_text) > 6000:
        doc_summary = full_text[:4000] + "\n...\n" + full_text[-2000:]
    else:
        doc_summary = full_text

    llm = ChatAnthropic(
        model=settings.VISION_MODEL,
        anthropic_api_key=settings.ANTHROPIC_API_KEY,
        temperature=0,
        max_tokens=150,
        max_retries=3,
    )

    for i, chunk in enumerate(chunks):
        if i > 0 and i % 5 == 0:
            await asyncio.sleep(1)  # Rate limit: pause every 5 chunks
        try:
            context = await llm.ainvoke(
                _CONTEXT_PROMPT.format(document=doc_summary, chunk=chunk.page_content)
            )
            chunk.page_content = f"[{context.content.strip()}]\n{chunk.page_content}"
        except Exception:
            logger.warning("Failed to generate context for chunk, skipping")

    return chunks


# ──────────────────────────────────────
# Common: URL extraction + metadata + upsert
# ──────────────────────────────────────

_URL_PATTERN = re.compile(r'https?://[^\s\)\]\>"\']+')


async def _upsert(
    chunks: list[Document],
    namespace: str,
    extra_metadata: dict[str, Any],
    full_text: str = "",
    skip_enrichment: bool = False,
) -> int:
    """Contextual retrieval → URL extraction → metadata → upsert to Pinecone."""
    if full_text and not skip_enrichment:
        chunks = await _enrich_with_context(chunks, full_text)

    for chunk in chunks:
        chunk.metadata.update(extra_metadata)
        urls = _URL_PATTERN.findall(chunk.page_content)
        if urls:
            chunk.metadata["urls"] = urls

    vectorstore = get_vectorstore(namespace)
    await vectorstore.aadd_documents(chunks)
    return len(chunks)
