"""Ingestion v2: Opus 4.7-first universal pipeline.

Replaces v1's 5 format-specific paths (PDF/DOCX/XLSX/legacy/markdown)
with one path: any file → PDF → Opus parse+chunk → reuse _upsert.

See docs/superpowers/specs/2026-04-18-ingest-v2-design.md for design.
"""
from __future__ import annotations

import base64
import io
import json
import logging
import os
import re
import zipfile
from functools import lru_cache
from typing import Any

from langchain_core.documents import Document
from langchain_core.messages import HumanMessage, SystemMessage

from ingest.services._v2_prompts import (
    SYSTEM_PROMPT,
    USER_PROMPT_TEMPLATE,
    format_sidecar,
)
from ingest.services.vision import _looks_like_refusal

logger = logging.getLogger(__name__)


# OCR fallback constants (used when ingest_v2 detects a pure-scan PDF)
OCR_MODEL = "claude-haiku-4-5-20251001"
OCR_CONCURRENCY = 4
OCR_DPI = 200
OCR_MAX_TOKENS_PER_PAGE = 4096
PURE_SCAN_TEXT_THRESHOLD = 0


def _extract_json_from_fence(text):
    """Extract a JSON object from a markdown fence, tolerating extra prose.

    Opus 4.7 reliably outputs JSON inside a ```json … ``` fence when
    the system prompt asks for it. We accept either a fenced code block or
    bare JSON with extra text by finding the outermost { … } span.

    Returns ``None`` on parse failure or empty input — caller logs and treats
    as 0 chunks.
    """
    if not text:
        return None
    m = _JSON_FENCE_RE.search(text)
    candidate = m.group(1) if m else text
    try:
        return json.loads(candidate)
    except json.JSONDecodeError:
        start = candidate.find("{")
        end = candidate.rfind("}")
        if start >= 0 and end > start:
            try:
                return json.loads(candidate[start:end + 1])
            except json.JSONDecodeError:
                return None
    return None


def _text_from_response(response):
    """Return the concatenated text-block content from a langchain AIMessage.

    Opus 4.7 with adaptive thinking returns ``content`` as a list of typed
    blocks (``thinking``, ``text``, …). Plain string content is also handled
    for compatibility with simpler responses (and tests).
    """
    content = response.content
    if isinstance(content, str):
        return content
    if isinstance(content, list):
        parts = []
        for b in content:
            if isinstance(b, dict) and b.get("type") == "text":
                parts.append(b.get("text", ""))
        return "\n".join(parts)
    return str(content or "")


def _pdf_to_image_blocks(pdf_bytes, dpi=150):
    """Rasterise every PDF page to PNG and wrap as Anthropic image content blocks.

    Used by ``opus_parse_and_chunk`` when the PDF has no text layer.
    Mirrors what claude.ai does internally for PDF uploads — Anthropic's PDF
    ``document`` block path is unreliable on long pure-scan Thai legal docs
    (silent ``chunks=[]`` even with an OCR sidecar). Sending pre-rendered
    images bypasses that path.

    DPI 150 balances OCR fidelity against payload size: a 24-page Thai legal
    scan at 150 DPI is ~36K input tokens vs ~5K for the document block.

    Anthropic limits a single request to 100 image blocks. This function
    raises ``ValueError`` if the PDF exceeds that.
    """
    import pymupdf
    blocks = []
    doc = pymupdf.open(stream=pdf_bytes, filetype="pdf")
    try:
        if doc.page_count > 100:
            raise ValueError(
                f"_pdf_to_image_blocks: {doc.page_count} pages exceeds "
                f"Anthropic's 100-image limit"
            )
        for page in doc:
            pix = page.get_pixmap(dpi=dpi, alpha=False)
            png = pix.tobytes("png")
            blocks.append({
                "type": "image",
                "source": {
                    "type": "base64",
                    "media_type": "image/png",
                    "data": base64.standard_b64encode(png).decode("utf-8"),
                },
            })
    finally:
        doc.close()
    return blocks


@lru_cache(maxsize=1)
def _get_ocr_client():
    """Cached raw AsyncAnthropic client for per-page vision OCR.

    Intentionally uses the raw anthropic SDK rather than langchain_anthropic:
    the OCR call is a single image + text block with no tool use, and raw
    SDK has cleaner async semantics and direct access to response bodies for
    debugging. Tests monkeypatch this function (or its return value) — the
    cache is small so a cache_clear() in a fixture is cheap.
    """
    from anthropic import AsyncAnthropic
    from shared.config import settings
    return AsyncAnthropic(api_key=settings.ANTHROPIC_API_KEY, max_retries=3)


OCR_PROMPT = (
    "สกัดข้อความทั้งหมดที่มองเห็นจากภาพสแกนหน้านี้ "
    "คงรูปโครงสร้าง (หัวข้อ ย่อหน้า รายการหัวข้อ ตาราง) เท่าที่ทำได้ "
    "- หัวข้อให้ขึ้นบรรทัดใหม่ "
    "- ตารางให้ใช้ pipe markdown | col1 | col2 | "
    "- ไม่ต้องใส่คำอธิบายใด ๆ หรือบอกว่าเป็นภาพอะไร ให้คืนเฉพาะข้อความ "
    "- ถ้ามีภาษาอังกฤษปนให้คงไว้ตามต้นฉบับ"
)


async def ocr_pdf_pages(pdf_bytes: bytes, filename: str) -> dict[int, str]:
    """Per-page vision OCR using Haiku 4.5, parallelized up to OCR_CONCURRENCY.

    Returns ``{1-based page: text}``. A per-page exception yields an empty
    string for that page (partial OCR is better than total fail). If every
    page raises, the function raises ``RuntimeError``.
    """
    import asyncio
    import pymupdf

    client = _get_ocr_client()
    doc = pymupdf.open(stream=pdf_bytes, filetype="pdf")
    n_pages = doc.page_count
    # Render page images up front so the async tasks don't fight over the pymupdf handle.
    page_pngs: dict[int, bytes] = {}
    try:
        for i, page in enumerate(doc):
            page_pngs[i + 1] = page.get_pixmap(dpi=OCR_DPI).tobytes("png")
    finally:
        doc.close()

    sem = asyncio.Semaphore(OCR_CONCURRENCY)

    async def _one(page_num: int) -> tuple[int, str | BaseException]:
        async with sem:
            try:
                b64 = base64.standard_b64encode(page_pngs[page_num]).decode("ascii")
                resp = await client.messages.create(
                    model=OCR_MODEL,
                    max_tokens=OCR_MAX_TOKENS_PER_PAGE,
                    messages=[{
                        "role": "user",
                        "content": [
                            {"type": "image", "source": {"type": "base64", "media_type": "image/png", "data": b64}},
                            {"type": "text", "text": OCR_PROMPT},
                        ],
                    }],
                )
                text_parts = [b.text for b in resp.content if getattr(b, "type", "") == "text"]
                return page_num, "\n".join(text_parts).strip()
            except (asyncio.CancelledError, Exception) as exc:
                return page_num, exc

    results = await asyncio.gather(*[_one(p) for p in range(1, n_pages + 1)])

    out: dict[int, str] = {}
    n_failed = 0
    for page_num, payload in results:
        if isinstance(payload, BaseException):
            logger.warning("ocr_pdf_pages(%s): page %d failed: %r", filename, page_num, payload)
            out[page_num] = ""
            n_failed += 1
        else:
            out[page_num] = payload

    if n_failed == n_pages:
        raise RuntimeError(f"OCR failed for all pages of {filename!r}")

    logger.info(
        "ocr_pdf_pages(%s): OCR complete — %d pages, %d chars total, %d failed",
        filename, n_pages, sum(len(v) for v in out.values()), n_failed,
    )
    return out


_OCR_DOCX_PAGE_HEADING_RE = re.compile(r"^หน้า\s+(\d+)\s*$")
_JSON_FENCE_RE = re.compile(r"```(?:json)?\s*(\{.*?\})\s*```", re.DOTALL)
_UNREADABLE_PAGE_RE = re.compile(r"\[page \d+: unreadable\]")


def _read_ocr_docx_as_pages(file_bytes: bytes) -> dict[int, str]:
    """Parse a .ocr.docx (produced by ``scripts/ocr_pdf_via_opus.py``) into
    ``{1-based page: joined text}`` matching ``ocr_pdf_pages``'s shape.

    The script writes: level-1 heading (doc title, ignored) → level-2
    ``หน้า N`` heading opening each page → paragraphs with OCR'd text.
    This parser walks paragraphs in document order; a Heading 2 matching
    ``r'^หน้า\\s+(\\d+)$'`` opens a new page bucket, subsequent non-heading
    paragraphs accumulate into it joined by ``\\n``.

    Fallback: if no page markers are found (hand-edited or tool-produced
    docx), return ``{1: <all paragraphs joined>}`` so ingest proceeds as
    single-page rather than fails outright.

    Raises ``ValueError`` when python-docx cannot open the bytes (wraps
    ``docx.opc.exceptions.PackageNotFoundError`` so callers see a stable
    exception type).
    """
    from docx import Document as DocxDocument
    from docx.opc.exceptions import PackageNotFoundError

    try:
        doc = DocxDocument(io.BytesIO(file_bytes))
    except (PackageNotFoundError, zipfile.BadZipFile, KeyError) as exc:
        raise ValueError(f"not a valid .ocr.docx: {exc}") from exc

    pages: dict[int, list[str]] = {}
    current_page: int | None = None

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style_name = (para.style.name or "") if para.style else ""
        if style_name == "Heading 2":
            m = _OCR_DOCX_PAGE_HEADING_RE.match(text)
            if m:
                current_page = int(m.group(1))
                pages.setdefault(current_page, [])
                continue
            # Other level-2 headings (unexpected) — fall through as content
            # on the currently open page, if any.
        if current_page is None:
            continue
        pages[current_page].append(text)

    if not pages:
        # No page markers — collect all non-empty paragraphs into page 1.
        joined = "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
        return {1: joined}

    return {p: "\n".join(lines) for p, lines in pages.items()}


def _format_pages_for_text_only(ocr_sidecar: dict[int, str]) -> str:
    """Render per-page OCR text for the text-only Opus prompt (no PDF block).

    Distinct from ``format_ocr_sidecar`` (in ``_v2_prompts``) which treats
    the rendered PDF image as ground truth and OCR as assistive. Here the
    text IS the ground truth — Opus has no image to cross-check. Page
    markers use ``### Page N`` to match ``format_ocr_sidecar``'s convention
    so chunk ``page`` attribution is consistent regardless of which path
    produced the text.
    """
    if not ocr_sidecar:
        return "(empty document)"
    lines: list[str] = []
    for page_num in sorted(ocr_sidecar.keys()):
        text = ocr_sidecar[page_num]
        lines.append(f"### Page {page_num}")
        lines.append(text if text else "(no text extracted on this page)")
        lines.append("")
    return "\n".join(lines).rstrip()


@lru_cache(maxsize=1)
def _get_opus_llm():
    """Return the Opus 4.7 LLM used for v2 parse+chunk (cached per process).

    Adaptive thinking is ENABLED. No tool binding — the system prompt
    instructs Opus to emit a ``\\`\\`\\`json`` fence and we parse it in Python.
    Earlier iterations of this pipeline used ``bind_tools`` + ``tool_choice``
    but empirically Opus would opt out of the tool call on long Thai legal
    docs (silent ``chunks=[]``). Free-form text + JSON fence reliably
    produces parseable output.

    Cached so a batch audit (e.g. 14 PDFs) does not re-instantiate the
    HTTP-pooled client once per file. Tests monkeypatch this function
    directly — ``monkeypatch.setattr`` replaces the module attribute so
    the cached original is bypassed entirely during the test.
    """
    from langchain_anthropic import ChatAnthropic
    from shared.config import settings
    return ChatAnthropic(
        model=settings.OCR_MODEL,
        anthropic_api_key=settings.ANTHROPIC_API_KEY,
        # 32K output budget: slide decks (45+ pages × ~500 tokens/chunk)
        # and dense announcement PDFs (23+ student records) both
        # overflowed the previous 8K cap, surfacing as silent 0-chunk
        # returns when the tool_call JSON was truncated mid-array.
        max_tokens=32000,
        max_retries=3,
        thinking={"type": "adaptive"},
    )


def ensure_pdf(file_bytes: bytes, filename: str) -> bytes:
    """Normalize any supported input to PDF bytes.

    PDF inputs are passed through untouched (byte-identity preserved).
    Other formats (DOC/DOCX/XLS/XLSX/PPT/PPTX/CSV) are converted via LibreOffice
    (delegated to the battle-tested `_convert_to_pdf` helper).
    """
    ext = os.path.splitext(filename)[1].lower()
    if ext == ".pdf":
        return file_bytes
    if ext in {".doc", ".docx", ".xls", ".xlsx", ".ppt", ".pptx", ".csv"}:
        from ingest.services.ingest_helpers import _convert_to_pdf
        return _convert_to_pdf(file_bytes, ext)
    raise ValueError(f"ensure_pdf: unsupported extension '{ext}' for '{filename}'")


def extract_hyperlinks(pdf_bytes: bytes) -> list[dict]:
    """Extract hidden hyperlink URIs from every page as sidecar metadata.

    Returns a list of ``{"page": int, "text": str, "uri": str}`` entries.
    ``text`` is the visible anchor text (from the PDF text layer clipped
    to the link rectangle); ``uri`` is the target URL. Page numbering is
    1-based to match how users refer to pages.

    URIs that already appear as visible text on the page are skipped —
    Opus 4.7 reads them directly from the rendered PDF and a duplicate
    sidecar entry would cause it to emit duplicate markdown links.

    Opus 4.7 reading a PDF cannot see link annotations — only the rendered
    visual. We inject this sidecar so generated chunks can still emit
    inline ``[anchor](uri)`` markdown links when relevant.
    """
    import pymupdf

    links: list[dict] = []
    doc = pymupdf.open(stream=pdf_bytes, filetype="pdf")
    try:
        for page_index, page in enumerate(doc):
            page_text = page.get_text("text")
            for link in page.get_links():
                uri = link.get("uri", "")
                if not uri:
                    continue
                # Skip URIs already visible in the rendered page text — Opus
                # will read them directly. Sidecar is for *hidden* link
                # annotations only (matches v1 behavior + v2 prompt contract).
                if uri in page_text:
                    continue
                rect = link.get("from", pymupdf.Rect())
                anchor_text = page.get_text("text", clip=rect).strip() or uri
                links.append({
                    "page": page_index + 1,
                    "text": anchor_text,
                    "uri": uri,
                })
    finally:
        doc.close()
    return links


def extract_page_text(pdf_bytes: bytes) -> dict[int, str]:
    """Return {1-based page index: extracted text layer} for the PDF.

    Used by ``ingest_v2`` to detect pure-scan PDFs (all values empty). Also
    exposes per-page text so future work can feed hybrid PDFs' native text
    to Opus as an OCR-equivalent sidecar without triggering the LLM path.
    """
    import pymupdf

    out: dict[int, str] = {}
    doc = pymupdf.open(stream=pdf_bytes, filetype="pdf")
    try:
        for i, page in enumerate(doc):
            out[i + 1] = page.get_text("text") or ""
    finally:
        doc.close()
    return out


async def opus_parse_and_chunk(
    pdf_bytes,
    hyperlinks,
    filename,
    mode,
):
    """Send PDF (or rasterised page images) to Opus 4.7, parse JSON-fence response.

    No tool binding. The system prompt instructs Opus to emit one
    ``\\`\\`\\`json`` fence containing ``{"chunks": [...]}``. We accept text output
    and parse the fence.

    ``mode`` selects the content shape:
    - ``"document"`` — text-layer PDF: ``[document_block, text_block]``
    - ``"images"``   — pure-scan PDF: ``[image_block, …, text_block]``

    Empty / refusal / [page N: unreadable] chunks are filtered out. Any failure
    to parse JSON returns ``[]`` (caller logs and treats as 0 chunks).
    """
    sidecar_block = format_sidecar(hyperlinks)
    user_text = USER_PROMPT_TEMPLATE.format(
        filename=filename,
        sidecar_block=sidecar_block,
    )

    if mode == "images":
        image_blocks = _pdf_to_image_blocks(pdf_bytes)
        human_content = [*image_blocks, {"type": "text", "text": user_text}]
    elif mode == "document":
        pdf_b64 = base64.standard_b64encode(pdf_bytes).decode("utf-8")
        human_content = [
            {
                "type": "document",
                "source": {
                    "type": "base64",
                    "media_type": "application/pdf",
                    "data": pdf_b64,
                },
            },
            {"type": "text", "text": user_text},
        ]
    else:
        raise ValueError(f"opus_parse_and_chunk: unknown mode {mode!r}")

    messages = [
        SystemMessage(content=SYSTEM_PROMPT),
        HumanMessage(content=human_content),
    ]
    response = await _get_opus_llm().ainvoke(messages)

    meta = response.response_metadata or {}
    stop_reason = meta.get("stop_reason")
    usage = meta.get("usage", {}) or {}
    text = _text_from_response(response)
    logger.info(
        "opus_parse(%s mode=%s): stop=%s in_tok=%s out_tok=%s text_len=%d",
        filename, mode, stop_reason,
        usage.get("input_tokens"), usage.get("output_tokens"), len(text),
    )

    parsed = _extract_json_from_fence(text)
    if not parsed or "chunks" not in parsed:
        logger.warning(
            "opus_parse(%s): could not parse JSON. preview=%r",
            filename, text[:800],
        )
        return []

    raw_chunks = parsed.get("chunks") or []
    cleaned = []
    for c in raw_chunks:
        t = (c.get("text") or "").strip()
        if not t:
            continue
        if _UNREADABLE_PAGE_RE.fullmatch(t):
            continue
        if _looks_like_refusal(t):
            logger.warning(
                "opus_parse(%s): dropping refusal chunk on page %s",
                filename, c.get("page", "?"),
            )
            continue
        try:
            page = int(c.get("page", 1))
        except (TypeError, ValueError):
            page = 1
        cleaned.append(Document(
            page_content=t,
            metadata={
                "page": page,
                "section_path": c.get("section_path", "") or "",
                "has_table": bool(c.get("has_table", False)),
            },
        ))
    return cleaned


async def ingest_v2(
    file_bytes: bytes,
    filename: str,
    namespace: str,
    tenant_id: str,
    doc_category: str = "general",
    url: str = "",
    download_link: str = "",
    drive_file_id: str = "",
) -> int:
    """v2 universal ingestion entrypoint.

    Pipeline:
      1. Route by filename:
         - ``.ocr.docx`` → ``_read_ocr_docx_as_pages`` → text-only Opus path
         - other → ``ensure_pdf`` + pure-scan detection → OCR or not → Opus
      2. ``opus_parse_and_chunk`` returns chunks with page/section_path/has_table
      3. ``_upsert`` (reused from v1) — Cohere embed, Pinecone atomic swap,
         BM25 cross-process invalidation

    ``drive_file_id`` (optional) stores the Drive file ID in chunk metadata
    so admin delete can remove the Drive file by ID even after rename —
    name-based lookup breaks when users rename files in Drive after ingest.
    """
    from ingest.services.ingest_helpers import _build_metadata, _upsert

    pdf_bytes: bytes | None
    hyperlinks: list[dict]
    ocr_sidecar: dict[int, str] | None

    if filename.lower().endswith(".ocr.docx"):
        # OCR'd content already — skip LibreOffice + PDF path entirely.
        ocr_sidecar = _read_ocr_docx_as_pages(file_bytes)
        pdf_bytes = None
        hyperlinks = []
        logger.info(
            "ingest_v2(%s): .ocr.docx detected — %d pages, %d total chars, skipping LibreOffice + PDF path",
            filename, len(ocr_sidecar), sum(len(t) for t in ocr_sidecar.values()),
        )
    else:
        pdf_bytes = ensure_pdf(file_bytes, filename)
        page_text = extract_page_text(pdf_bytes)
        total_text_chars = sum(len(t) for t in page_text.values())
        if total_text_chars <= PURE_SCAN_TEXT_THRESHOLD:
            logger.info(
                "ingest_v2(%s): pure-scan detected (0 text chars across %d pages), running OCR",
                filename, len(page_text),
            )
            ocr_sidecar = await ocr_pdf_pages(pdf_bytes, filename)
        else:
            logger.debug(
                "ingest_v2(%s): text-layer PDF (%d chars across %d pages), OCR skipped",
                filename, total_text_chars, len(page_text),
            )
            ocr_sidecar = None
        hyperlinks = extract_hyperlinks(pdf_bytes)

    chunks = await opus_parse_and_chunk(
        pdf_bytes, hyperlinks, filename, ocr_sidecar=ocr_sidecar,
    )

    if not chunks:
        logger.warning("ingest_v2(%s): Opus returned 0 chunks — nothing upserted", filename)
        return 0

    metadata = _build_metadata(
        tenant_id=tenant_id,
        source_type="pdf",  # v2 universalizes to pdf
        source_filename=filename,
        doc_category=doc_category,
        url=url,
        download_link=download_link,
        drive_file_id=drive_file_id,
    )
    return await _upsert(chunks, namespace, metadata)
