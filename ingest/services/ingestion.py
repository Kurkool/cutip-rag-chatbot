"""Document ingestion pipeline: PDF, DOCX, XLSX/CSV, and web content to Pinecone vector store."""

import asyncio
import io
import logging
import os
import re
import subprocess
import tempfile
import time
from typing import Any

import pandas as pd
import pymupdf
from docx import Document as DocxDocument
from langchain_core.documents import Document

from fastapi import HTTPException

from shared.config import settings
from shared.services import usage
from shared.services.resilience import call_with_backoff
from shared.services.vectorstore import (
    PINECONE_PAGE_SIZE,
    fetch_metadata_batch,
    get_raw_index,
    get_vectorstore,
    list_all_vector_ids,
)
from ingest.services.chunking import _smart_chunk, _chunk_pages
from ingest.services.enrichment import _enrich_with_context
from ingest.services.vision import interpret_spreadsheet, parse_page_image

logger = logging.getLogger(__name__)

# ──────────────────────────────────────
# Vision / batch-size constants
# ──────────────────────────────────────

_VISION_THRESHOLD = settings.PDF_VISION_THRESHOLD
_PDF_BATCH_SIZE = settings.PDF_BATCH_SIZE
_XLSX_BATCH_ROWS = settings.XLSX_BATCH_ROWS


# ──────────────────────────────────────
# Metadata
# ──────────────────────────────────────

def _build_metadata(
    tenant_id: str,
    source_type: str,
    source_filename: str = "",
    doc_category: str = "general",
    url: str = "",
    download_link: str = "",
) -> dict[str, Any]:
    return {
        "tenant_id": tenant_id,
        "source_type": source_type,
        "source_filename": source_filename,
        "doc_category": doc_category,
        "url": url,
        "download_link": download_link,
    }


# ──────────────────────────────────────
# Legacy format conversion (.doc, .xls, .ppt → PDF)
# ──────────────────────────────────────

def _convert_to_pdf(file_bytes: bytes, src_ext: str) -> bytes:
    """Convert legacy formats to PDF using LibreOffice headless."""
    with tempfile.TemporaryDirectory() as tmpdir:
        src_path = os.path.join(tmpdir, f"input{src_ext}")
        with open(src_path, "wb") as f:
            f.write(file_bytes)

        result = subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", tmpdir, src_path],
            capture_output=True,
            timeout=settings.LIBREOFFICE_TIMEOUT,
        )
        if result.returncode != 0:
            logger.error("LibreOffice failed: %s", result.stderr.decode(errors="replace"))

        pdf_path = os.path.join(tmpdir, "input.pdf")
        if not os.path.exists(pdf_path):
            raise RuntimeError(f"LibreOffice conversion failed for {src_ext}")

        with open(pdf_path, "rb") as f:
            return f.read()


async def ingest_legacy(
    file_bytes: bytes,
    filename: str,
    namespace: str,
    tenant_id: str,
    skip_enrichment: bool = False,
    doc_category: str = "general",
    url: str = "",
    download_link: str = "",
) -> int:
    """Convert .doc/.xls/.ppt → PDF → Vision pipeline."""
    ext = os.path.splitext(filename)[1].lower()
    pdf_bytes = _convert_to_pdf(file_bytes, ext)
    return await ingest_pdf(
        file_bytes=pdf_bytes,
        filename=filename,
        namespace=namespace,
        tenant_id=tenant_id,
        skip_enrichment=skip_enrichment,
        doc_category=doc_category,
        url=url,
        download_link=download_link,
    )


# ──────────────────────────────────────
# Duplicate Detection: delete old vectors before re-ingest
# ──────────────────────────────────────

def _delete_existing_vectors(
    namespace: str,
    source_filename: str,
    older_than_ts: float | None = None,
) -> int:
    """Delete vectors matching ``source_filename`` in the namespace.

    When ``older_than_ts`` is supplied, only vectors whose ``ingest_ts``
    metadata is strictly less than it are deleted — used by ``_upsert`` to
    dedup OLD copies after a fresh upsert has already landed (atomic swap:
    new vectors survive even if the delete is interrupted). When omitted,
    every matching vector is deleted (legacy behaviour retained for
    administrative wipes).

    Pinecone serverless does not support metadata-filter delete, so we list
    IDs + fetch metadata + delete by ID in batches.
    """
    ids = list_all_vector_ids(namespace)
    if not ids:
        return 0

    index = get_raw_index()
    deleted = 0
    for i in range(0, len(ids), PINECONE_PAGE_SIZE):
        batch = ids[i:i + PINECONE_PAGE_SIZE]
        fetched = index.fetch(ids=batch, namespace=namespace)
        matching: list[str] = []
        for vid, vec in fetched.vectors.items():
            meta = vec.metadata or {}
            if meta.get("source_filename") != source_filename:
                continue
            if older_than_ts is not None:
                # Keep vectors at or newer than this timestamp (the ones we
                # just upserted). Missing/zero ingest_ts counts as old.
                vec_ts = meta.get("ingest_ts", 0) or 0
                if vec_ts >= older_than_ts:
                    continue
            matching.append(vid)
        if matching:
            index.delete(ids=matching, namespace=namespace)
            deleted += len(matching)

    if deleted:
        logger.info(
            "Dedup: deleted %d old vectors for '%s' in namespace '%s'",
            deleted, source_filename, namespace,
        )
    return deleted


# ──────────────────────────────────────
# PDF Ingestion (Hybrid: text extraction + Vision fallback)
# ──────────────────────────────────────

async def ingest_pdf(
    file_bytes: bytes,
    filename: str,
    namespace: str,
    tenant_id: str,
    skip_enrichment: bool = False,
    doc_category: str = "general",
    url: str = "",
    download_link: str = "",
) -> int:
    """
    Hybrid PDF ingestion:
    - Pages with enough text → PyMuPDF text extraction (fast, free)
    - Pages with little/no text (tables, forms, scanned) → Claude Vision (accurate)
    - Hidden hyperlinks extracted from all pages

    Handles: text, forms, slides, scanned, tables, diagrams, hidden hyperlinks.

    Old vectors for the same filename are deleted INSIDE ``_upsert`` only
    AFTER the new vectors land — see atomic-swap note in ``_upsert``.
    """
    tmp_path = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name

        doc = pymupdf.open(tmp_path)
        # Pre-flight guards. Encrypted PDFs and oversized PDFs both need doc
        # closed on the way out — use an explicit try/finally anchor so no
        # exit path leaks the PyMuPDF native handle.
        preflight_ok = False
        try:
            if doc.is_encrypted:
                logger.error("Password-protected PDF: %s", filename)
                raise HTTPException(
                    status_code=400,
                    detail="Password-protected PDF — please remove encryption before uploading.",
                )

            # Reject oversized PDFs with a clear message rather than letting
            # them burn through Vision quota + time out.
            if doc.page_count > settings.PDF_MAX_PAGES:
                raise HTTPException(
                    status_code=413,
                    detail=(
                        f"PDF has {doc.page_count} pages (limit: {settings.PDF_MAX_PAGES}). "
                        f"Please split into smaller files."
                    ),
                )
            preflight_ok = True
        finally:
            if not preflight_ok:
                doc.close()

        # Phase 1: Extract text + classify each page
        pages_data = []
        vision_pages = []
        for i, page in enumerate(doc):
            text = page.get_text("text").strip()
            hidden_links = _extract_hyperlinks(page)
            page_num = i + 1

            # Has tables? (PyMuPDF can detect table structures)
            has_tables = bool(page.find_tables().tables)

            if len(text) >= _VISION_THRESHOLD and not has_tables:
                # Enough text, no tables → use text extraction (fast)
                if hidden_links:
                    text += "\n\n**Links in this page:**\n" + "\n".join(
                        f"- [{t}]({u})" for t, u in hidden_links
                    )
                pages_data.append({"text": text, "page": page_num})
            else:
                # Low text or has tables → need Vision
                pix = page.get_pixmap(dpi=150)
                vision_pages.append({
                    "img_bytes": pix.tobytes("png"),
                    "links": hidden_links,
                    "page_num": page_num,
                })
        doc.close()

        # Phase 2: Process Vision pages in batches
        logger.info(
            "%s: %d text pages, %d vision pages",
            filename, len(pages_data), len(vision_pages),
        )
        for batch_start in range(0, len(vision_pages), _PDF_BATCH_SIZE):
            if batch_start > 0:
                await asyncio.sleep(2)  # Rate limit: pause between batches
            batch = vision_pages[batch_start:batch_start + _PDF_BATCH_SIZE]
            tasks = [_process_pdf_page(p) for p in batch]
            results = await asyncio.gather(*tasks, return_exceptions=True)
            for result in results:
                if isinstance(result, dict) and result.get("text"):
                    pages_data.append(result)

        # Track vision calls
        if vision_pages:
            await usage.track(tenant_id, "vision_call", len(vision_pages))

        # Sort by page number
        pages_data.sort(key=lambda p: p["page"])
        page_texts = pages_data

        if not page_texts:
            return 0

        full_text = "\n\n---\n\n".join(p["text"] for p in page_texts)

        # Detect slides vs document
        avg_chars = len(full_text) / max(len(page_texts), 1)
        is_slides = len(page_texts) > 5 and avg_chars < 500

        if is_slides:
            chunks = _chunk_pages(page_texts, source=filename)
        else:
            chunks = _smart_chunk(full_text, source=filename)
            for chunk in chunks:
                if "page" not in chunk.metadata:
                    chunk.metadata["page"] = 1

        metadata = _build_metadata(
            tenant_id, "pdf", filename, doc_category, url, download_link
        )
        return await _upsert(chunks, namespace, metadata, full_text=full_text, skip_enrichment=skip_enrichment)
    finally:
        if tmp_path and os.path.exists(tmp_path):
            os.unlink(tmp_path)


async def _process_pdf_page(page_data: dict) -> dict:
    """Process a single PDF page: Vision + hyperlinks."""
    try:
        markdown = await parse_page_image(page_data["img_bytes"])
        if page_data["links"]:
            markdown += "\n\n**Links in this page:**\n" + "\n".join(
                f"- [{text}]({uri})" for text, uri in page_data["links"]
            )
        return {"text": markdown, "page": page_data["page_num"]}
    except Exception:
        logger.warning("Failed to process page %d", page_data["page_num"])
        return {"text": "", "page": page_data["page_num"]}


def _extract_hyperlinks(page) -> list[tuple[str, str]]:
    """Extract hidden hyperlinks from a PDF page (not visible as text)."""
    page_text = page.get_text("text")
    links = []
    for link in page.get_links():
        uri = link.get("uri", "")
        if not uri or uri in page_text:
            continue  # Skip if URI is already visible in text
        rect = link.get("from", pymupdf.Rect())
        text = page.get_text("text", clip=rect).strip() or uri
        links.append((text, uri))
    return links


# ──────────────────────────────────────
# DOCX Ingestion (paragraphs + tables + images)
# ──────────────────────────────────────

async def ingest_docx(
    file_bytes: bytes,
    filename: str,
    namespace: str,
    tenant_id: str,
    skip_enrichment: bool = False,
    doc_category: str = "general",
    url: str = "",
    download_link: str = "",
) -> int:
    """
    DOCX → extract paragraphs + tables + images → combined markdown → chunk → embed.
    No more missing tables or images. Dedup happens atomically inside ``_upsert``.
    """
    doc = DocxDocument(io.BytesIO(file_bytes))

    # Pre-flight: count images before processing so we can reject oversized
    # docs rather than blowing through Vision quota mid-stream.
    image_blobs = [
        rel.target_part.blob
        for rel in doc.part.rels.values()
        if "image" in rel.reltype
    ]
    if len(image_blobs) > settings.DOCX_MAX_IMAGES:
        raise HTTPException(
            status_code=413,
            detail=(
                f"DOCX has {len(image_blobs)} images (limit: {settings.DOCX_MAX_IMAGES}). "
                f"Please split into smaller files."
            ),
        )

    parts = []

    # 1. Extract paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            # Detect heading styles
            if para.style and para.style.name.startswith("Heading"):
                level = para.style.name.replace("Heading ", "")
                try:
                    hashes = "#" * int(level)
                except ValueError:
                    hashes = "##"
                parts.append(f"{hashes} {text}")
            else:
                parts.append(text)

    # 2. Extract tables → markdown tables
    for table in doc.tables:
        md_table = _docx_table_to_markdown(table)
        if md_table:
            parts.append(md_table)

    # 3. Extract images → Claude Vision (parallel, bounded, backed-off).
    # Previously this loop was sequential and would 429 after ~10-15 images
    # on a heavy DOCX. Now INGEST_CONCURRENCY calls in flight + retry.
    if image_blobs:
        semaphore = asyncio.Semaphore(settings.INGEST_CONCURRENCY)
        caption_tasks = [
            call_with_backoff(
                lambda img=blob: parse_page_image(img),
                semaphore=semaphore,
                max_retries=settings.INGEST_MAX_RETRIES,
                label=f"docx_image[{i}]",
            )
            for i, blob in enumerate(image_blobs)
        ]
        captions = await asyncio.gather(*caption_tasks)
        success_count = 0
        for caption in captions:
            if caption:
                parts.append(f"[Image content: {caption}]")
                success_count += 1
        await usage.track(tenant_id, "vision_call", success_count)
        if success_count < len(image_blobs):
            logger.warning(
                "DOCX '%s': %d/%d images processed successfully",
                filename, success_count, len(image_blobs),
            )

    full_text = "\n\n".join(parts)
    chunks = _smart_chunk(full_text, source=filename)
    metadata = _build_metadata(
        tenant_id, "docx", filename, doc_category, url, download_link
    )
    return await _upsert(chunks, namespace, metadata, full_text=full_text, skip_enrichment=skip_enrichment)


def _docx_table_to_markdown(table) -> str:
    """Convert a python-docx table to markdown table."""
    rows = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        rows.append(cells)

    if not rows:
        return ""

    # First row as header
    header = "| " + " | ".join(rows[0]) + " |"
    separator = "| " + " | ".join("---" for _ in rows[0]) + " |"
    body = "\n".join(
        "| " + " | ".join(row) + " |" for row in rows[1:]
    )
    return f"{header}\n{separator}\n{body}"


# ──────────────────────────────────────
# Markdown Ingestion (web content)
# ──────────────────────────────────────

async def ingest_markdown(
    content: str,
    title: str,
    namespace: str,
    tenant_id: str,
    skip_enrichment: bool = False,
    doc_category: str = "general",
    url: str = "",
    download_link: str = "",
) -> int:
    """Markdown content (from Jina Reader) → smart chunk → embed.

    Dedup happens atomically inside ``_upsert``.
    """
    source_name = title or url or "web_content"
    chunks = _smart_chunk(content, source=source_name)
    metadata = _build_metadata(
        tenant_id, "web", source_name, doc_category, url, download_link
    )
    return await _upsert(chunks, namespace, metadata, full_text=content, skip_enrichment=skip_enrichment)


# ──────────────────────────────────────
# Spreadsheet Ingestion (Claude interprets, batched for large sheets)
# ──────────────────────────────────────

async def ingest_spreadsheet(
    file_bytes: bytes,
    filename: str,
    namespace: str,
    tenant_id: str,
    skip_enrichment: bool = False,
    doc_category: str = "general",
    url: str = "",
    download_link: str = "",
) -> tuple[int, int]:
    """
    XLSX/CSV → raw dump → Claude interprets (batched for large sheets)
    → markdown → chunk → embed. Dedup happens atomically inside ``_upsert``.
    """
    is_csv = filename.lower().endswith(".csv")
    total_chunks = 0

    # Pre-flight: reject oversized spreadsheets before processing. We count
    # total rows across all sheets; CSVs are always one sheet.
    if is_csv:
        df_preflight = pd.read_csv(io.BytesIO(file_bytes), header=None)
        total_rows_preflight = len(df_preflight.dropna(how="all"))
    else:
        xls_preflight = pd.ExcelFile(io.BytesIO(file_bytes))
        total_rows_preflight = sum(
            len(xls_preflight.parse(s, header=None).dropna(how="all"))
            for s in xls_preflight.sheet_names
        )
    if total_rows_preflight > settings.XLSX_MAX_ROWS:
        raise HTTPException(
            status_code=413,
            detail=(
                f"Spreadsheet has {total_rows_preflight} data rows "
                f"(limit: {settings.XLSX_MAX_ROWS}). Please split into smaller files."
            ),
        )

    if is_csv:
        df = pd.read_csv(io.BytesIO(file_bytes), header=None)
        structured, api_calls = await _interpret_dataframe(df)
        chunks = _smart_chunk(structured, source=filename)
        metadata = _build_metadata(
            tenant_id, "spreadsheet", filename, doc_category, url, download_link
        )
        total_chunks = await _upsert(chunks, namespace, metadata, full_text=structured, skip_enrichment=skip_enrichment)
        await usage.track(tenant_id, "vision_call", api_calls)
        return 1, total_chunks

    # XLSX with multiple sheets: collect all chunks FIRST, then upsert ONCE.
    # Previously each sheet called _upsert separately, and each call's
    # post-upsert dedup (by source_filename + older_than_ts) wiped out the
    # previous sheet's freshly-upserted chunks because they shared the same
    # base filename but had an older ingest_ts. Regression caught 2026-04-17
    # on `ตารางเรียน ปี 2568 ปโท CU-TIP.xlsx` — 16 chunks reported, 1 survived.
    xls = pd.ExcelFile(io.BytesIO(file_bytes))
    all_chunks: list[Document] = []
    all_text_parts: list[str] = []
    sheets_processed = 0
    vision_calls = 0
    for sheet_name in xls.sheet_names:
        df = xls.parse(sheet_name, header=None)
        if df.dropna(how="all").empty:
            continue
        structured, api_calls = await _interpret_dataframe(df, sheet_name=sheet_name)
        vision_calls += api_calls
        chunks = _smart_chunk(structured, source=f"{filename} - {sheet_name}")
        all_chunks.extend(chunks)
        all_text_parts.append(structured)
        sheets_processed += 1

    if vision_calls:
        await usage.track(tenant_id, "vision_call", vision_calls)

    if not all_chunks:
        return 0, 0

    # Single atomic upsert across all sheets — one ingest_ts stamp, one dedup.
    full_text = "\n\n".join(all_text_parts)
    metadata = _build_metadata(
        tenant_id, "spreadsheet", filename, doc_category, url, download_link
    )
    total_chunks = await _upsert(
        all_chunks, namespace, metadata, full_text=full_text, skip_enrichment=skip_enrichment,
    )
    return sheets_processed, total_chunks


async def _interpret_dataframe(df: pd.DataFrame, sheet_name: str = "") -> tuple[str, int]:
    """Send DataFrame to Claude in batches for large sheets. Returns (text, api_call_count).

    Large sheets fan out to Claude in parallel (bounded by INGEST_CONCURRENCY
    + exponential backoff on rate limits). Previously batches ran sequentially,
    which for a 10k-row sheet meant ~5 minutes wall-clock; parallel execution
    cuts that to under a minute while staying within rate limits.
    """
    df_clean = df.dropna(how="all").dropna(axis=1, how="all")
    total_rows = len(df_clean)

    if total_rows <= _XLSX_BATCH_ROWS:
        # Small sheet: send all at once
        raw = _raw_dataframe_dump(df_clean)
        prefix = f"Sheet: {sheet_name}\n\n" if sheet_name else ""
        return await interpret_spreadsheet(f"{prefix}{raw}"), 1

    # Large sheet: build per-batch prompts, then fan out.
    prompts: list[tuple[int, str]] = []
    for start in range(0, total_rows, _XLSX_BATCH_ROWS):
        batch = df_clean.iloc[start:start + _XLSX_BATCH_ROWS]
        raw = _raw_dataframe_dump(batch)
        prefix = f"Sheet: {sheet_name} (rows {start+1}-{start+len(batch)})\n\n"
        prompts.append((start, f"{prefix}{raw}"))

    semaphore = asyncio.Semaphore(settings.INGEST_CONCURRENCY)
    tasks = [
        call_with_backoff(
            lambda p=prompt: interpret_spreadsheet(p),
            semaphore=semaphore,
            max_retries=settings.INGEST_MAX_RETRIES,
            label=f"xlsx_batch[{sheet_name}:{start}]",
        )
        for start, prompt in prompts
    ]
    results = await asyncio.gather(*tasks)

    # Preserve original order; drop failed batches (already logged by helper).
    parts = [r for r in results if r]
    if len(parts) < len(results):
        logger.warning(
            "Spreadsheet '%s': %d/%d batches interpreted successfully",
            sheet_name or "csv", len(parts), len(results),
        )
    return "\n\n".join(parts), len(parts)


def _raw_dataframe_dump(df: pd.DataFrame) -> str:
    """Dump raw cell data as readable text for Claude to interpret."""
    lines = []
    for i, row in df.iterrows():
        values = []
        for j, val in enumerate(row):
            if pd.notna(val):
                values.append(f"[{j}]={val}")
        if values:
            lines.append(f"Row {i}: {' | '.join(values)}")
    return "\n".join(lines)


# ──────────────────────────────────────
# Common: URL extraction + metadata + upsert
# ──────────────────────────────────────

_URL_PATTERN = re.compile(r'https?://[^\s\)\]\>"\']+')


async def _upsert(
    chunks: list[Document],
    namespace: str,
    extra_metadata: dict[str, Any],
    full_text: str = "",
    skip_enrichment: bool = False,
) -> int:
    """Contextual retrieval → URL extraction → metadata → upsert to Pinecone.

    Atomic swap: new vectors are upserted FIRST; only after they land do we
    delete older copies of the same ``source_filename``. If anything before
    or during upsert fails, old vectors survive and the doc remains
    searchable. Previously delete-first meant a failed upsert = data loss
    (e.g. ``slide.pdf`` vanishing during a Cloud Run timeout on 2026-04-17).
    """
    if full_text and not skip_enrichment:
        chunks = await _enrich_with_context(chunks, full_text)

    # Stamp each chunk with a monotonically-increasing ingest timestamp so
    # dedup can distinguish "just-upserted" vectors from prior generations.
    ingest_ts = time.time()
    for chunk in chunks:
        chunk.metadata.update(extra_metadata)
        chunk.metadata["ingest_ts"] = ingest_ts
        urls = _URL_PATTERN.findall(chunk.page_content)
        if urls:
            # Pinecone caps metadata at 40KB per vector. A page with many
            # footnote URLs can exceed that and cause a silent upsert failure.
            # Cap at 20 URLs (well below the limit for typical URL lengths).
            chunk.metadata["urls"] = urls[:20]

    vectorstore = get_vectorstore(namespace)
    await vectorstore.aadd_documents(chunks)

    # NEW vectors are now live. Dedup OLD copies of the same source_filename
    # (strictly older ingest_ts). If this step fails or is interrupted, the
    # worst case is a brief window of duplicate results — not data loss.
    source_filename = extra_metadata.get("source_filename", "")
    if source_filename:
        await asyncio.to_thread(
            _delete_existing_vectors,
            namespace, source_filename, older_than_ts=ingest_ts,
        )

    # Two-tier BM25 invalidation:
    #   1. Local (same-process) — drops this worker's cache if it happens to
    #      hold a stale copy. No-op on first ingest after cold start.
    from shared.services.bm25_cache import invalidate_bm25_cache
    invalidate_bm25_cache(namespace)

    #   2. Cross-process — bump tenant.bm25_invalidate_ts so chat-api
    #      (running in a DIFFERENT container) notices on next request and
    #      re-warms its own BM25 cache. Without this, chat keeps stale
    #      results until the container cycles (hours/days).
    tenant_id = extra_metadata.get("tenant_id", "")
    if tenant_id:
        from shared.services import firestore as firestore_service
        try:
            await firestore_service.bump_bm25_invalidate_ts(tenant_id)
        except Exception:
            logger.warning(
                "Failed to bump bm25_invalidate_ts for tenant %s", tenant_id,
                exc_info=True,
            )

    return len(chunks)
