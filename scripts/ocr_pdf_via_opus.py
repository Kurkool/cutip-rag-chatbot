"""OCR a pure-scan PDF via Opus 4.7 vision, emit a .docx v2 can ingest.

Use when `diag_pdf.py` reports 0 text chars across all pages (pure image scan
with no text layer). v2's Opus parse+chunk step empirically returns 0 chunks
on long pure-scan PDFs; feeding a .docx (which v2 converts via LibreOffice
with fonts-thai-tlwg) bypasses that failure mode because the converted PDF
carries a real text layer.

Output: `<input>.ocr.docx` next to the input PDF.

Run:
    PYTHONPATH=. .venv/Scripts/python.exe scripts/ocr_pdf_via_opus.py <path>
"""
from __future__ import annotations

import argparse
import asyncio
import base64
import os
import subprocess
import sys
import time
from pathlib import Path

sys.stdout.reconfigure(encoding="utf-8")


def _load_anthropic_key_from_secret_manager() -> None:
    try:
        key = subprocess.check_output(
            ["gcloud", "secrets", "versions", "access", "latest", "--secret=ANTHROPIC_API_KEY"],
            shell=True,
        ).decode().strip()
        if key:
            os.environ["ANTHROPIC_API_KEY"] = key
            print(f"ocr_pdf: loaded ANTHROPIC_API_KEY from Secret Manager ({len(key)} chars)")
    except Exception as exc:
        print(f"ocr_pdf: WARNING — could not fetch ANTHROPIC_API_KEY from Secret Manager ({exc}); falling back to env/.env")


_load_anthropic_key_from_secret_manager()

import pymupdf
from anthropic import AsyncAnthropic
from docx import Document

OCR_PROMPT = (
    "สกัดข้อความทั้งหมดที่มองเห็นจากภาพสแกนหน้านี้ "
    "คงรูปโครงสร้าง (หัวข้อ ย่อหน้า รายการหัวข้อ ตาราง) เท่าที่ทำได้ "
    "- หัวข้อให้ขึ้นบรรทัดใหม่ "
    "- ตารางให้ใช้ pipe markdown | col1 | col2 | "
    "- ไม่ต้องใส่คำอธิบายใด ๆ หรือบอกว่าเป็นภาพอะไร ให้คืนเฉพาะข้อความ "
    "- ถ้ามีภาษาอังกฤษปนให้คงไว้ตามต้นฉบับ"
)

MODEL = "claude-opus-4-7"
DPI = 200
CONCURRENCY = 8
MAX_TOKENS = 4096


async def ocr_page(client: AsyncAnthropic, pdf_path: Path, page_num: int, sem: asyncio.Semaphore) -> tuple[int, str]:
    async with sem:
        doc = pymupdf.open(pdf_path)
        page = doc[page_num]
        png_bytes = page.get_pixmap(dpi=DPI).tobytes("png")
        doc.close()
        img_b64 = base64.standard_b64encode(png_bytes).decode("ascii")

        t0 = time.monotonic()
        resp = await client.messages.create(
            model=MODEL,
            max_tokens=MAX_TOKENS,
            messages=[{
                "role": "user",
                "content": [
                    {
                        "type": "image",
                        "source": {"type": "base64", "media_type": "image/png", "data": img_b64},
                    },
                    {"type": "text", "text": OCR_PROMPT},
                ],
            }],
        )
        dur = time.monotonic() - t0
        text_parts = [b.text for b in resp.content if getattr(b, "type", "") == "text"]
        text = "\n".join(text_parts).strip()
        print(f"  page {page_num + 1:>3}: {len(text):>5} chars in {dur:5.1f}s")
        return page_num, text


async def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("pdf_path", type=Path)
    parser.add_argument("--out", type=Path, default=None, help="output .docx path (default: <input>.ocr.docx)")
    args = parser.parse_args()

    pdf_path: Path = args.pdf_path
    out_path: Path = args.out or pdf_path.with_suffix(".ocr.docx")

    if not pdf_path.exists():
        sys.exit(f"not found: {pdf_path}")

    doc = pymupdf.open(pdf_path)
    n_pages = doc.page_count
    doc.close()

    print(f"ocr_pdf: input   {pdf_path} ({n_pages} pages)")
    print(f"ocr_pdf: model   {MODEL}  dpi={DPI}  concurrency={CONCURRENCY}")
    print(f"ocr_pdf: output  {out_path}")

    client = AsyncAnthropic()
    sem = asyncio.Semaphore(CONCURRENCY)

    t_all = time.monotonic()
    tasks = [ocr_page(client, pdf_path, i, sem) for i in range(n_pages)]
    results = await asyncio.gather(*tasks, return_exceptions=True)

    ok: list[tuple[int, str]] = []
    errors: list[tuple[int, BaseException]] = []
    for r in results:
        if isinstance(r, BaseException):
            errors.append((-1, r))
        else:
            ok.append(r)
    ok.sort(key=lambda x: x[0])
    total_dur = time.monotonic() - t_all

    docx = Document()
    docx.add_heading(pdf_path.stem, level=1)
    for page_num, text in ok:
        docx.add_heading(f"หน้า {page_num + 1}", level=2)
        for para in text.split("\n"):
            if para.strip():
                docx.add_paragraph(para)
    docx.save(out_path)

    total_chars = sum(len(t) for _, t in ok)
    print()
    print(f"ocr_pdf: done in {total_dur:.1f}s — {total_chars:,} chars across {len(ok)} pages, {len(errors)} errors")
    print(f"ocr_pdf: wrote  {out_path}  ({out_path.stat().st_size:,} bytes)")
    if errors:
        for _, e in errors:
            print(f"  error: {e!r}")


if __name__ == "__main__":
    asyncio.run(main())
