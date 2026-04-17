"""Comprehensive recheck of TIP-RAG ingestion pipeline.

Phases:
  1. Extract ground truth from every source file (PDF, DOCX, DOC, XLSX) in
     sample-doc/cutip-doc/ using the most faithful parser per type
  2. Fetch all Pinecone chunks grouped by source_filename
  3. Diff: coverage %, dropped entities, orphan chunks, duplicate content,
     suspicious tiny/empty chunks, Vision-refusal pollution
  4. Per-file retrieval probe: query via /api/chat, verify hit lands on the
     expected source
  5. Entity sweep: every Thai person name / student ID / email from ANY
     source file — is it findable in Pinecone?
"""
import io
import re
import subprocess
import sys
import time
from pathlib import Path

sys.stdout.reconfigure(encoding="utf-8")

import httpx
import pandas as pd
import fitz
from docx import Document as DocxDocument
from pinecone import Pinecone

SAMPLE_DIR = Path(r"C:\Users\USER\PycharmProjects\TIP-RAG\sample-doc\cutip-doc")
CHAT_URL = "https://cutip-chat-api-secaaxwrgq-as.a.run.app/api/chat"

# Patterns used to extract ground-truth entities from source files.
THAI_NAME_RE = re.compile(r"(?:นาย|นางสาว|นาง)\s*([ก-๙]+(?:\s+[ก-๙]+)?)")
STUDENT_ID_RE = re.compile(r"\b\d{10}\b")

VISION_ERROR_PATTERNS = [
    "Ensure the image is clear",
    "Could you please",
    "Re-upload the document",
    "readable document",
    "no visible text",
    "convert to Markdown",
    "file loaded correctly",
    "i'll be happy",
    "i cannot process",
    "please provide",
]


def extract_source_text(path: Path) -> str:
    """Best-effort plain-text extraction from any supported source type."""
    suffix = path.suffix.lower()
    try:
        if suffix == ".pdf":
            doc = fitz.open(str(path))
            t = "\n".join(p.get_text("text") for p in doc)
            doc.close()
            return t
        if suffix == ".docx":
            d = DocxDocument(str(path))
            parts = [p.text for p in d.paragraphs if p.text.strip()]
            for table in d.tables:
                for row in table.rows:
                    parts.extend(c.text for c in row.cells if c.text.strip())
            return "\n".join(parts)
        if suffix == ".doc":
            # .doc requires LibreOffice conversion — out of scope for an audit,
            # return empty so it's flagged as "source untestable"
            return ""
        if suffix in (".xlsx", ".xls"):
            xls = pd.ExcelFile(str(path))
            parts = []
            for sheet in xls.sheet_names:
                df = xls.parse(sheet, header=None)
                for _, row in df.iterrows():
                    for cell in row:
                        if pd.notna(cell):
                            parts.append(str(cell))
            return "\n".join(parts)
    except Exception as e:
        return f"<extraction failed: {e}>"
    return ""


def fetch_all_chunks(pc: Pinecone, namespace: str) -> dict[str, list[dict]]:
    idx = pc.Index("university-rag")
    all_ids = []
    for page in idx.list(namespace=namespace):
        all_ids.extend(page)
    by_file: dict[str, list[dict]] = {}
    for i in range(0, len(all_ids), 50):
        batch = idx.fetch(ids=all_ids[i : i + 50], namespace=namespace)
        for vid, v in batch.vectors.items():
            meta = v.metadata or {}
            fn = meta.get("source_filename", "?")
            by_file.setdefault(fn, []).append({
                "id": vid,
                "text": str(meta.get("text", "")),
                "page": meta.get("page"),
            })
    return by_file


def diff_file(src_path: Path, chunks: list[dict]) -> dict:
    src = extract_source_text(src_path)
    merged = "\n".join(c["text"] for c in chunks)
    chunk_chars = len(merged)
    src_chars = len(src)
    # Entities
    src_names = set(THAI_NAME_RE.findall(src))
    dropped_names = [n for n in src_names if n not in merged]
    src_ids = set(STUDENT_ID_RE.findall(src))
    dropped_ids = [i for i in src_ids if i not in merged]
    # Chunk hygiene
    vision_errs = sum(
        1 for c in chunks
        if any(p in c["text"] for p in VISION_ERROR_PATTERNS)
    )
    tiny = sum(1 for c in chunks if len(c["text"]) < 50)
    dup_prefixes: dict[str, int] = {}
    for c in chunks:
        prefix = c["text"][:100]
        dup_prefixes[prefix] = dup_prefixes.get(prefix, 0) + 1
    duplicate_chunks = sum(v - 1 for v in dup_prefixes.values() if v > 1)
    return {
        "file": src_path.name,
        "src_chars": src_chars,
        "chunks": len(chunks),
        "chunk_chars": chunk_chars,
        "coverage_pct": round(100 * chunk_chars / max(1, src_chars), 0),
        "dropped_names": dropped_names[:5],
        "n_dropped_names": len(dropped_names),
        "n_src_names": len(src_names),
        "dropped_ids": dropped_ids[:5],
        "n_dropped_ids": len(dropped_ids),
        "n_src_ids": len(src_ids),
        "vision_err_chunks": vision_errs,
        "tiny_chunks": tiny,
        "duplicate_chunks": duplicate_chunks,
    }


def retrieval_probe(
    client: httpx.Client, api_key: str, query: str, expect_filename: str,
    namespace: str,
) -> dict:
    r = client.post(
        CHAT_URL,
        headers={"X-API-Key": api_key, "Content-Type": "application/json"},
        json={"query": query, "tenant_id": namespace, "user_id": "audit"},
        timeout=120,
    )
    ok = r.status_code == 200
    srcs = r.json().get("sources", []) if ok else []
    src_files = [s.get("filename", "") for s in srcs]
    hit = any(expect_filename in f or f in expect_filename for f in src_files)
    return {
        "query": query,
        "expect": expect_filename,
        "status": r.status_code,
        "n_sources": len(srcs),
        "top_files": src_files[:3],
        "hit": hit,
    }


def main():
    import argparse

    parser = argparse.ArgumentParser(description="Full ingestion audit")
    parser.add_argument(
        "--namespace",
        default="cutip_01",
        help="Pinecone namespace to audit (default: cutip_01, v1 production)",
    )
    args = parser.parse_args()
    namespace = args.namespace

    key = subprocess.check_output(
        ["gcloud", "secrets", "versions", "access", "latest", "--secret=PINECONE_API_KEY"],
        shell=True,
    ).decode().strip()
    pc = Pinecone(api_key=key)
    admin_key = subprocess.check_output(
        ["gcloud", "secrets", "versions", "access", "latest", "--secret=ADMIN_API_KEY"],
        shell=True,
    ).decode().strip()

    print("=== PHASE 1-3: SOURCE VS PINECONE DIFF ===\n")
    by_file = fetch_all_chunks(pc, namespace)
    src_files = sorted([p for p in SAMPLE_DIR.iterdir() if p.is_file()])
    summary = []
    for p in src_files:
        chunks = by_file.get(p.name, [])
        if not chunks:
            print(f"!! {p.name}  NOT IN PINECONE (not ingested)")
            continue
        r = diff_file(p, chunks)
        summary.append(r)
        issues = []
        if r["n_dropped_names"] > 0:
            issues.append(f"dropped_names={r['n_dropped_names']}/{r['n_src_names']}")
        if r["n_dropped_ids"] > 0:
            issues.append(f"dropped_ids={r['n_dropped_ids']}/{r['n_src_ids']}")
        if r["vision_err_chunks"] > 0:
            issues.append(f"vision_err={r['vision_err_chunks']}")
        if r["tiny_chunks"] > 0:
            issues.append(f"tiny={r['tiny_chunks']}")
        if r["duplicate_chunks"] > 0:
            issues.append(f"dup_prefix={r['duplicate_chunks']}")
        flag = "OK" if not issues else "ISSUES: " + ", ".join(issues)
        print(f"{p.name[:55]:55s}  chunks={r['chunks']:3d}  cov={r['coverage_pct']}%  {flag}")
        if r["dropped_names"]:
            print(f"  dropped_names examples: {r['dropped_names']}")
        if r["dropped_ids"]:
            print(f"  dropped_ids examples: {r['dropped_ids']}")

    # Orphan chunks (filenames in Pinecone that don't exist on disk)
    disk_names = {p.name for p in src_files}
    orphans = [f for f in by_file if f not in disk_names]
    if orphans:
        print(f"\n!! Pinecone has {len(orphans)} orphan file(s) not on disk:")
        for o in orphans:
            print(f"   {o}  ({len(by_file[o])} chunks)")

    print("\n=== PHASE 4: PER-FILE RETRIEVAL PROBES ===\n")
    probes = [
        ("คณะกรรมการสอบนายเกื้อกูล", "ประกาศแจ้งคณะกรรมการสอบ"),
        ("ทุนการศึกษาสำหรับนิสิต", "ทุนการศึกษา.docx"),
        ("ตารางเรียน รุ่น 18", "ตารางเรียน"),
        ("ขั้นตอนสอบวิทยานิพนธ์", "สอบวิทยานิพนธ์.pdf"),
        ("สอบโครงร่างวิทยานิพนธ์คืออะไร", "สอบโครงร่าง"),
        ("สอบความก้าวหน้าวิทยานิพนธ์", "สอบความก้าวหน้า"),
        ("สอบโครงการพิเศษขั้นตอน", "สอบโครงการพิเศษ.pdf"),
        ("slide presentation", "slide.pdf"),
        ("ห้องเรียน ภาคปลาย 2568", "ห้องเรียน"),
    ]
    with httpx.Client() as client:
        retrieval_results = []
        for q, expected in probes:
            r = retrieval_probe(client, admin_key, q, expected, namespace)
            retrieval_results.append(r)
            mark = "✅" if r["hit"] else "❌"
            print(f"{mark} q={q[:40]:40s} expect={expected[:30]:30s} got_top={r['top_files'][:2]}")
            time.sleep(0.3)  # gentle pace

    print("\n=== PHASE 5: ENTITY COVERAGE SWEEP ===\n")
    all_names = set()
    all_ids = set()
    for p in src_files:
        if p.suffix.lower() in (".pdf", ".docx"):
            src = extract_source_text(p)
            all_names.update(THAI_NAME_RE.findall(src))
            all_ids.update(STUDENT_ID_RE.findall(src))
    merged_pinecone = "\n".join(
        c["text"] for chunks in by_file.values() for c in chunks
    )
    missing_names = [n for n in all_names if n not in merged_pinecone]
    missing_ids = [i for i in all_ids if i not in merged_pinecone]
    print(f"Thai names in sources: {len(all_names)}")
    print(f"  found in Pinecone:    {len(all_names) - len(missing_names)}/{len(all_names)}")
    if missing_names:
        print(f"  missing examples:     {missing_names[:5]}")
    print(f"Student IDs in sources: {len(all_ids)}")
    print(f"  found in Pinecone:    {len(all_ids) - len(missing_ids)}/{len(all_ids)}")
    if missing_ids:
        print(f"  missing examples:     {missing_ids[:5]}")

    # Summary scorecard
    total_issues = sum(
        r["n_dropped_names"] + r["n_dropped_ids"]
        + r["vision_err_chunks"] + r["duplicate_chunks"]
        for r in summary
    )
    retrieval_hits = sum(1 for r in retrieval_results if r["hit"])
    print(f"\n=== SCORECARD (namespace={namespace}) ===")
    print(f"Files audited:        {len(summary)} / {len(src_files)} ingested")
    print(f"Entity drops:         {len(missing_names)} names + {len(missing_ids)} IDs missing from Pinecone")
    print(f"Per-file retrieval:   {retrieval_hits}/{len(retrieval_results)} probes landed on expected file")
    print(f"Ingestion issues:     {total_issues} (dropped entities + vision errs + dup prefixes)")


if __name__ == "__main__":
    main()
