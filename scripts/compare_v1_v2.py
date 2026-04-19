"""Side-by-side Pinecone-level quality comparison of v1 (cutip_01) vs v2 (cutip_v2_audit).

No /api/chat calls — no Anthropic/Cohere cost. Pure Pinecone fetch + local source diff.
"""
import re
import subprocess
import sys
from collections import Counter
from pathlib import Path

sys.stdout.reconfigure(encoding="utf-8")

import fitz
import pandas as pd
from docx import Document as DocxDocument
from pinecone import Pinecone

SAMPLE_DIR = Path(r"C:\Users\USER\PycharmProjects\TIP-RAG\sample-doc\cutip-doc")
NS_V1 = "cutip_01"
NS_V2 = "cutip_v2_audit"

THAI_NAME_RE = re.compile(r"(?:นาย|นางสาว|นาง)\s*([ก-๙]+(?:\s+[ก-๙]+)?)")
STUDENT_ID_RE = re.compile(r"\b\d{10}\b")
EMAIL_RE = re.compile(r"[\w.+-]+@[\w-]+\.[\w.-]+")

VISION_ERRS = [
    "Ensure the image is clear", "Could you please", "Re-upload the document",
    "readable document", "no visible text", "convert to Markdown",
    "file loaded correctly", "i'll be happy", "i cannot process", "please provide",
]


def extract_source_text(path: Path) -> str:
    suffix = path.suffix.lower()
    try:
        if suffix == ".pdf":
            doc = fitz.open(str(path))
            t = "\n".join(p.get_text("text") for p in doc)
            doc.close()
            return t
        if suffix == ".docx":
            d = DocxDocument(str(path))
            parts = [p.text for p in d.paragraphs if p.text.strip()]
            for table in d.tables:
                for row in table.rows:
                    parts.extend(c.text for c in row.cells if c.text.strip())
            return "\n".join(parts)
        if suffix in (".xlsx", ".xls"):
            xls = pd.ExcelFile(str(path))
            parts = []
            for sheet in xls.sheet_names:
                df = xls.parse(sheet, header=None)
                for _, row in df.iterrows():
                    for cell in row:
                        if pd.notna(cell):
                            parts.append(str(cell))
            return "\n".join(parts)
    except Exception as e:
        return f"<extraction failed: {e}>"
    return ""


def fetch_all(pc: Pinecone, namespace: str):
    idx = pc.Index("university-rag")
    ids = []
    for page in idx.list(namespace=namespace):
        ids.extend(page)
    by_file: dict[str, list[dict]] = {}
    all_meta_keys: Counter = Counter()
    for i in range(0, len(ids), 50):
        batch = idx.fetch(ids=ids[i:i + 50], namespace=namespace)
        for vid, v in batch.vectors.items():
            meta = v.metadata or {}
            all_meta_keys.update(meta.keys())
            fn = meta.get("source_filename", "?")
            by_file.setdefault(fn, []).append({
                "id": vid,
                "text": str(meta.get("text", "")),
                "meta": meta,
            })
    return by_file, all_meta_keys, len(ids)


def analyze_file(src: str, chunks: list[dict]) -> dict:
    merged = "\n".join(c["text"] for c in chunks)
    src_names = set(THAI_NAME_RE.findall(src))
    src_ids = set(STUDENT_ID_RE.findall(src))
    src_emails = set(EMAIL_RE.findall(src))
    return {
        "chunks": len(chunks),
        "chunk_chars": len(merged),
        "src_chars": len(src),
        "coverage": round(100 * len(merged) / max(1, len(src)), 1),
        "avg_chunk": int(len(merged) / max(1, len(chunks))),
        "n_src_names": len(src_names),
        "found_names": sum(1 for n in src_names if n in merged),
        "n_src_ids": len(src_ids),
        "found_ids": sum(1 for i in src_ids if i in merged),
        "n_src_emails": len(src_emails),
        "found_emails": sum(1 for e in src_emails if e in merged),
        "tiny": sum(1 for c in chunks if len(c["text"]) < 50),
        "vision_err": sum(1 for c in chunks if any(p in c["text"] for p in VISION_ERRS)),
        "dup_prefix": sum(
            v - 1 for v in Counter(c["text"][:100] for c in chunks).values() if v > 1
        ),
    }


def main():
    key = subprocess.check_output(
        ["gcloud", "secrets", "versions", "access", "latest", "--secret=PINECONE_API_KEY"],
        shell=True,
    ).decode().strip()
    pc = Pinecone(api_key=key)

    print(f"\n{'='*100}")
    print(f"{'FETCHING CHUNKS':^100}")
    print(f"{'='*100}\n")
    v1, v1_keys, v1_total = fetch_all(pc, NS_V1)
    v2, v2_keys, v2_total = fetch_all(pc, NS_V2)
    print(f"v1 ({NS_V1:20s}) total vectors: {v1_total:4d}  files: {len(v1):3d}")
    print(f"v2 ({NS_V2:20s}) total vectors: {v2_total:4d}  files: {len(v2):3d}")

    print(f"\n{'='*100}")
    print(f"{'METADATA FIELDS (keys present in at least one chunk)':^100}")
    print(f"{'='*100}")
    all_keys = sorted(set(v1_keys.keys()) | set(v2_keys.keys()))
    print(f"{'field':35s} {'v1 count':>10s} {'v2 count':>10s}")
    for k in all_keys:
        print(f"{k:35s} {v1_keys.get(k, 0):>10d} {v2_keys.get(k, 0):>10d}")

    print(f"\n{'='*100}")
    print(f"{'PER-FILE COMPARISON':^100}")
    print(f"{'='*100}\n")
    src_files = sorted([p for p in SAMPLE_DIR.iterdir() if p.is_file()])
    header = f"{'file':42s} | {'v1_ch':>5s} {'v1_cov':>6s} {'v1_ent':>8s} | {'v2_ch':>5s} {'v2_cov':>6s} {'v2_ent':>8s}"
    print(header)
    print("-" * len(header))

    aggregates = {"v1": {}, "v2": {}}
    per_file = []
    for p in src_files:
        src = extract_source_text(p)
        v1c = v1.get(p.name, [])
        v2c = v2.get(p.name, [])
        a1 = analyze_file(src, v1c) if v1c else None
        a2 = analyze_file(src, v2c) if v2c else None

        def fmt(a):
            if not a:
                return (f"{'-':>5s}", f"{'-':>6s}", f"{'-':>8s}")
            total_ent = a["n_src_names"] + a["n_src_ids"] + a["n_src_emails"]
            found_ent = a["found_names"] + a["found_ids"] + a["found_emails"]
            return (
                f"{a['chunks']:>5d}",
                f"{a['coverage']:>5.0f}%",
                f"{found_ent}/{total_ent:>2d}".rjust(8),
            )

        c1, cov1, e1 = fmt(a1)
        c2, cov2, e2 = fmt(a2)
        print(f"{p.name[:42]:42s} | {c1} {cov1} {e1} | {c2} {cov2} {e2}")
        per_file.append({"file": p.name, "v1": a1, "v2": a2})

    print(f"\n{'='*100}")
    print(f"{'AGGREGATE SCORECARD':^100}")
    print(f"{'='*100}\n")

    for label, data in [("v1 (cutip_01)", v1), ("v2 (cutip_v2_audit)", v2)]:
        total_chunks = sum(len(cs) for cs in data.values())
        tiny = sum(1 for cs in data.values() for c in cs if len(c["text"]) < 50)
        vision = sum(1 for cs in data.values() for c in cs if any(p in c["text"] for p in VISION_ERRS))
        empty_meta = sum(1 for cs in data.values() for c in cs if not c["meta"].get("text"))
        avg_size = (
            sum(len(c["text"]) for cs in data.values() for c in cs) / max(1, total_chunks)
        )
        print(f"--- {label} ---")
        print(f"  total chunks:          {total_chunks}")
        print(f"  avg chunk size (chars): {avg_size:.0f}")
        print(f"  tiny (<50 char):       {tiny}")
        print(f"  vision-refusal chunks: {vision}")
        print(f"  empty-text chunks:     {empty_meta}")

    print(f"\n{'='*100}")
    print(f"{'ENTITY COVERAGE (pooled over ALL source files)':^100}")
    print(f"{'='*100}\n")
    all_names, all_ids, all_emails = set(), set(), set()
    for p in src_files:
        if p.suffix.lower() in (".pdf", ".docx", ".xlsx"):
            src = extract_source_text(p)
            all_names.update(THAI_NAME_RE.findall(src))
            all_ids.update(STUDENT_ID_RE.findall(src))
            all_emails.update(EMAIL_RE.findall(src))

    def pooled_merged(data):
        return "\n".join(c["text"] for cs in data.values() for c in cs)

    m1, m2 = pooled_merged(v1), pooled_merged(v2)
    print(f"{'entity type':20s} {'total':>7s} {'v1 found':>10s} {'v2 found':>10s}")
    print(f"{'Thai names':20s} {len(all_names):>7d} "
          f"{sum(1 for n in all_names if n in m1):>10d} "
          f"{sum(1 for n in all_names if n in m2):>10d}")
    print(f"{'Student IDs':20s} {len(all_ids):>7d} "
          f"{sum(1 for i in all_ids if i in m1):>10d} "
          f"{sum(1 for i in all_ids if i in m2):>10d}")
    print(f"{'Emails':20s} {len(all_emails):>7d} "
          f"{sum(1 for e in all_emails if e in m1):>10d} "
          f"{sum(1 for e in all_emails if e in m2):>10d}")

    # Orphans
    disk_names = {p.name for p in src_files}
    print(f"\n{'='*100}")
    print(f"{'ORPHAN / EXTRA FILES IN PINECONE':^100}")
    print(f"{'='*100}")
    for label, data in [("v1", v1), ("v2", v2)]:
        orphans = [f for f in data if f not in disk_names]
        if orphans:
            print(f"\n{label} orphans (not in sample-doc/cutip-doc/):")
            for o in orphans:
                print(f"  {o}  ({len(data[o])} chunks)")


if __name__ == "__main__":
    main()
