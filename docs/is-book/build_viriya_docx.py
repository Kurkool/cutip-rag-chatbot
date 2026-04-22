"""
Build VIRIYA IS manuscript to .docx with Thai formatting (TH Sarabun New + Chula margins).

Run from cutip-rag-chatbot/ dir:
    .venv/Scripts/python.exe docs/is-book/build_viriya_docx.py
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


BASE = Path(__file__).parent
MANUSCRIPT = BASE / "manuscript"
OUT = BASE / "VIRIYA-IS.docx"

CHAPTER_ORDER = [
    "frontmatter.md",
    "ch01-introduction.md",
    "ch02-literature-review.md",
    "ch03-methodology.md",
    "ch04-results.md",
    "ch05-business-feasibility.md",
    "ch06-financial-feasibility.md",
    "ch07-conclusion.md",
    "backmatter.md",
]

FONT = "TH Sarabun New"


def set_font_run(run, size_pt: int = 16, bold: bool = False):
    run.font.name = FONT
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    rPr = run._element.get_or_add_rPr()
    # Set fonts for ascii/hAnsi/cs/eastAsia
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), FONT)
    # Complex script size (half-points)
    szCs = rPr.find(qn("w:szCs"))
    if szCs is None:
        szCs = OxmlElement("w:szCs")
        rPr.append(szCs)
    szCs.set(qn("w:val"), str(size_pt * 2))
    # Thai language tag
    lang = rPr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        rPr.append(lang)
    lang.set(qn("w:val"), "th-TH")
    lang.set(qn("w:bidi"), "th-TH")


def apply_inline(paragraph, text: str, size_pt: int = 16):
    """Parse **bold** and `code` inline in a line of text."""
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_font_run(run, size_pt=size_pt, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_font_run(run, size_pt=size_pt)
            run.font.name = "Consolas"  # monospace for code
        else:
            run = paragraph.add_run(part)
            set_font_run(run, size_pt=size_pt)


def add_heading_para(doc, text: str, level: int, chapter_break: bool = False):
    """Add heading paragraph. level 1 = chapter (largest), 2 = section, etc."""
    para = doc.add_paragraph()
    # Page break before chapter-level heading
    if chapter_break:
        para.runs  # trigger init
        br_run = para.add_run()
        br_run.add_break(WD_BREAK.PAGE)
    sizes = {1: 22, 2: 18, 3: 16, 4: 16, 5: 16}
    size_pt = sizes.get(level, 16)
    if level == 1:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    set_font_run(run, size_pt=size_pt, bold=True)
    return para


def parse_markdown_table(lines: list[str], start_idx: int) -> tuple[list[list[str]], int]:
    """Parse markdown table starting at start_idx. Returns (rows, next_idx)."""
    rows = []
    i = start_idx
    while i < len(lines) and lines[i].lstrip().startswith("|"):
        row_raw = lines[i].strip()
        # Skip separator row (e.g., |---|---|)
        if re.match(r"^\|[\s\-:|]+\|?$", row_raw):
            i += 1
            continue
        # Parse cells
        cells = [c.strip() for c in row_raw.strip("|").split("|")]
        rows.append(cells)
        i += 1
    return rows, i


def add_table(doc, rows: list[list[str]]):
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx in range(n_cols):
            cell = table.cell(r_idx, c_idx)
            cell.text = ""  # clear default paragraph
            cell_text = row[c_idx] if c_idx < len(row) else ""
            para = cell.paragraphs[0]
            is_header = r_idx == 0
            apply_inline(para, cell_text, size_pt=14)
            if is_header:
                for run in para.runs:
                    run.font.bold = True
    doc.add_paragraph()  # spacing after


def process_file(doc, path: Path, is_chapter: bool):
    """Process one markdown file into doc."""
    with path.open(encoding="utf-8") as f:
        content = f.read()

    lines = content.split("\n")
    i = 0
    first_heading = True
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            i += 1
            continue

        # HTML comments → skip
        if stripped.startswith("<!--") and stripped.endswith("-->"):
            i += 1
            continue

        # HTML block (e.g., <div>) → paragraph
        if stripped.startswith("<") and stripped.endswith(">"):
            i += 1
            continue

        # Horizontal rule → page break
        if stripped == "---":
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            i += 1
            continue

        # Headings
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            chapter_break = is_chapter and first_heading and level == 1
            add_heading_para(doc, text, level, chapter_break=chapter_break)
            first_heading = False
            i += 1
            continue

        # Tables
        if stripped.startswith("|"):
            rows, next_i = parse_markdown_table(lines, i)
            add_table(doc, rows)
            i = next_i
            continue

        # Lists (simple bullet)
        if re.match(r"^[-*]\s+", stripped):
            para = doc.add_paragraph(style="List Bullet")
            content_text = re.sub(r"^[-*]\s+", "", stripped)
            apply_inline(para, content_text)
            # Set paragraph font
            for run in para.runs:
                set_font_run(run, size_pt=16)
            i += 1
            continue

        # Numbered list
        if re.match(r"^\d+\.\s+", stripped):
            para = doc.add_paragraph(style="List Number")
            content_text = re.sub(r"^\d+\.\s+", "", stripped)
            apply_inline(para, content_text)
            for run in para.runs:
                set_font_run(run, size_pt=16)
            i += 1
            continue

        # Blockquote
        if stripped.startswith(">"):
            para = doc.add_paragraph(style="Intense Quote")
            content_text = re.sub(r"^>\s*", "", stripped)
            apply_inline(para, content_text)
            i += 1
            continue

        # Regular paragraph — may span multiple lines until next heading/table/blank
        para_lines = [stripped]
        j = i + 1
        while j < len(lines):
            nxt = lines[j].strip()
            if not nxt:
                break
            if re.match(r"^#{1,6}\s", nxt):
                break
            if nxt.startswith("|"):
                break
            if nxt.startswith("---"):
                break
            if re.match(r"^[-*]\s", nxt) or re.match(r"^\d+\.\s", nxt):
                break
            para_lines.append(nxt)
            j += 1
        merged = " ".join(para_lines)
        para = doc.add_paragraph()
        apply_inline(para, merged)
        i = j


def main():
    doc = Document()

    # Global page setup per Chula thesis handbook 2548
    for section in doc.sections:
        section.left_margin = Inches(1.5)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

    # Default Normal style → TH Sarabun New 16pt
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(16)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts")) or OxmlElement("w:rFonts")
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), FONT)
    if rFonts not in list(rPr):
        rPr.append(rFonts)
    szCs = rPr.find(qn("w:szCs")) or OxmlElement("w:szCs")
    szCs.set(qn("w:val"), "32")
    if szCs not in list(rPr):
        rPr.append(szCs)
    lang = rPr.find(qn("w:lang")) or OxmlElement("w:lang")
    lang.set(qn("w:val"), "th-TH")
    lang.set(qn("w:bidi"), "th-TH")
    if lang not in list(rPr):
        rPr.append(lang)

    # Process files in order
    for idx, filename in enumerate(CHAPTER_ORDER):
        path = MANUSCRIPT / filename
        if not path.exists():
            print(f"  WARN: missing {filename}")
            continue
        is_chapter = idx > 0  # frontmatter (idx=0) doesn't need forced page break
        print(f"  processing {filename}...")
        process_file(doc, path, is_chapter=is_chapter)

    doc.save(str(OUT))
    print(f"\nSaved: {OUT}")
    print(f"File size: {OUT.stat().st_size:,} bytes")


if __name__ == "__main__":
    main()
