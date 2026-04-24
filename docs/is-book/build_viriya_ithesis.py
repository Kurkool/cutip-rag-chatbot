"""
Build VIRIYA IS manuscript into the iThesis Next Gen template.

Strategy: open VIRIYA-iThesis-Template.docx (iThesis-generated template) as base, find the
empty placeholder section between "คำอธิบายสัญลักษณ์" and "บรรณานุกรม", and inject
chapter content there using iThesis Heading styles. Also fill bibliography +
appendices from markdown.

User should still fill cover/approval/abstract/acknowledgement via the iThesis
add-in UI in Word (those have specific form-linked fields that the add-in
populates from student metadata).

Run from cutip-rag-chatbot/ dir:
    .venv/Scripts/python.exe docs/is-book/build_viriya_ithesis.py
"""
from __future__ import annotations

import re
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from lxml import etree


BASE = Path(__file__).parent
MANUSCRIPT = BASE / "manuscript"
TEMPLATE = BASE / "VIRIYA-iThesis-Template.docx"
FIGURES_DIR = BASE / "figures" / "generated"
USER_FIGURES_DIR = BASE / "figures" / "user-captured"
WATERMARK_IMG = BASE / "background-img.jpg"

# XML namespaces used by wp:anchor watermark
_NS_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_NS_WP = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
_NS_A = "http://schemas.openxmlformats.org/drawingml/2006/main"
_NS_R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

# Map figure number (as string e.g. "2.1") → PNG filename.
# A "user:" prefix resolves from figures/user-captured/ (user must save manually);
# otherwise resolve from figures/generated/ (auto-generated).
FIGURE_MAP = {
    "2.1": "fig_2_1_conceptual_framework.png",
    "3.1": "fig_3_1_methodology_phases.png",
    "3.2": "fig_3_2_system_architecture_new.png",
    "4.2": "fig_4_2_detailed_architecture_new.png",
    "4.3": "fig_4_3_evolution_timeline.png",
    "4.4": "user:fig_4_4_line_qa_good.png",
    "4.5": "user:fig_4_5_line_qa_gap.png",
    "4.6": "user:fig_4_6_admin_dashboard.png",
    "4.7": "user:fig_4_7_admin_chat_logs.png",
    "4.8": "user:fig_4_8_admin_upload.png",
    "4.9": "user:fig_4_9_admin_analytics.png",
}


def resolve_figure_path(filename: str) -> Path:
    """Resolve a FIGURE_MAP value to an actual path. 'user:name.png' → user-captured/;
    plain 'name.png' → generated/."""
    if filename.startswith("user:"):
        return USER_FIGURES_DIR / filename.removeprefix("user:")
    return FIGURES_DIR / filename
import sys as _sys
OUT_PRIMARY = BASE / "VIRIYA-IS-staging.docx"
OUT_FALLBACK = BASE / "VIRIYA-IS-staging-v2.docx"
try:
    # If primary is open in Word, writing will fail — use fallback
    _test = open(OUT_PRIMARY, "a+b")
    _test.close()
    OUT = OUT_PRIMARY
except (PermissionError, OSError):
    OUT = OUT_FALLBACK
    print(f"(PRIMARY {OUT_PRIMARY.name} locked — saving to {OUT.name} instead)", file=_sys.stderr)

# Order of chapter files to insert
CHAPTER_FILES = [
    "ch01-introduction.md",
    "ch02-literature-review.md",
    "ch03-methodology.md",
    "ch04-results.md",
    "ch05-business-feasibility.md",
    "ch06-financial-feasibility.md",
    "ch07-conclusion.md",
]


THAI_FONT = "TH Sarabun New"
BODY_SIZE_PT = 16  # Chula standard body text size


def force_thai_font_on_style(doc, style_name: str, *, size_pt: int | None = None,
                             bold: bool = False):
    """Override rFonts on a style to force TH Sarabun New (ascii/cs/hAnsi) and
    optionally set font size + bold. Also removes the theme accent color so
    headings render in default black instead of Word 2024's blue accent1
    (#0F4761). Fixes iThesis template's styles that inherit Aptos + blue via
    theme."""
    try:
        style = doc.styles[style_name]
    except KeyError:
        return
    rPr = style.element.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        style.element.insert(0, rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    # Remove theme attributes so explicit font takes effect
    for theme_attr in ("w:asciiTheme", "w:hAnsiTheme", "w:cstheme", "w:eastAsiaTheme"):
        if rFonts.get(qn(theme_attr)) is not None:
            del rFonts.attrib[qn(theme_attr)]
    rFonts.set(qn("w:ascii"), THAI_FONT)
    rFonts.set(qn("w:hAnsi"), THAI_FONT)
    rFonts.set(qn("w:cs"), THAI_FONT)
    # Remove theme-inherited color (accent1 renders as blue in Word 2024)
    color = rPr.find(qn("w:color"))
    if color is not None:
        rPr.remove(color)
    # Remove italic (iThesis template sets <w:i/> on Heading 4/6/8)
    for tag in ("w:i", "w:iCs"):
        el = rPr.find(qn(tag))
        if el is not None:
            rPr.remove(el)
    # Font size (w:sz + w:szCs use half-points: 16pt → 32)
    if size_pt is not None:
        half_pts = str(int(size_pt * 2))
        for tag in ("w:sz", "w:szCs"):
            el = rPr.find(qn(tag))
            if el is None:
                el = OxmlElement(tag)
                rPr.append(el)
            el.set(qn("w:val"), half_pts)
    # Bold (for headings — needed when all heading sizes match body size,
    # so bold provides visual hierarchy)
    if bold:
        for tag in ("w:b", "w:bCs"):
            if rPr.find(qn(tag)) is None:
                rPr.append(OxmlElement(tag))
    # Ensure Thai language tag — set w:val (primary language for Latin runs)
    # and w:bidi (complex script language). DO NOT set w:eastAsia — that slot
    # is for Chinese/Japanese/Korean; setting it to th-TH misleads Word into
    # applying CJK break/distribution rules on Thai text, which stretches
    # characters apart when distributed-justified.
    lang = rPr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        rPr.append(lang)
    lang.set(qn("w:val"), "th-TH")
    lang.set(qn("w:bidi"), "th-TH")
    # Clean up any stale eastAsia="th-TH" from earlier attempts
    if lang.get(qn("w:eastAsia")) is not None:
        del lang.attrib[qn("w:eastAsia")]


def _ensure_pPr(p_element):
    """Get or create <w:pPr> as the first child of a <w:p>."""
    pPr = p_element.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        p_element.insert(0, pPr)
    return pPr


def style_add_pPr_properties(doc, style_name: str, *, alignment: str | None = None,
                             first_line_indent_twips: int | None = None,
                             page_break_before: bool = False):
    """Add paragraph-level properties (alignment, first-line indent, page
    break before) to a style."""
    try:
        style = doc.styles[style_name]
    except KeyError:
        return
    pPr = style.element.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        # pPr must come before rPr in style definition
        rPr = style.element.find(qn("w:rPr"))
        if rPr is not None:
            style.element.insert(list(style.element).index(rPr), pPr)
        else:
            style.element.insert(0, pPr)
    if alignment is not None:
        jc = pPr.find(qn("w:jc"))
        if jc is None:
            jc = OxmlElement("w:jc")
            pPr.append(jc)
        jc.set(qn("w:val"), alignment)
    if first_line_indent_twips is not None:
        ind = pPr.find(qn("w:ind"))
        if ind is None:
            ind = OxmlElement("w:ind")
            pPr.append(ind)
        ind.set(qn("w:firstLine"), str(first_line_indent_twips))
    if page_break_before:
        pbb = pPr.find(qn("w:pageBreakBefore"))
        if pbb is None:
            pbb = OxmlElement("w:pageBreakBefore")
            pPr.append(pbb)


def apply_thai_lang_to_runs(p_element):
    """Force <w:lang w:val=th-TH w:bidi=th-TH w:eastAsia=th-TH/> on every run's
    rPr inside this paragraph. Without explicit run-level lang, Word auto-
    detects from Latin characters in the run and tags the run as English,
    which makes Thai Distributed alignment fall back to ASCII whitespace
    segmentation. That leaves half-empty lines and huge gaps when Thai words
    contain English tokens (e.g. '{Gao, 2024 #8}', 'Software-as-a-Service').
    With explicit Thai lang on every run, Word uses its Thai ICU segmenter
    and distributes cleanly."""
    for run in p_element.iter(qn("w:r")):
        rPr = run.find(qn("w:rPr"))
        if rPr is None:
            rPr = OxmlElement("w:rPr")
            run.insert(0, rPr)
        lang = rPr.find(qn("w:lang"))
        if lang is None:
            lang = OxmlElement("w:lang")
            rPr.append(lang)
        lang.set(qn("w:val"), "th-TH")
        lang.set(qn("w:bidi"), "th-TH")
        # Clean up any stale eastAsia
        if lang.get(qn("w:eastAsia")) is not None:
            del lang.attrib[qn("w:eastAsia")]


def apply_thai_layout_to_pPr(p_element):
    """Set wordWrap=0 so Word can break inside Thai phrases (needed for
    distributed to fill lines), but rely on document-level
    useAsianBreakRules compat flag to keep grapheme clusters intact."""
    pPr = _ensure_pPr(p_element)
    # Clean up any leftover flags from earlier experiments
    for tag in ("w:autoSpaceDE", "w:autoSpaceDN", "w:bidi"):
        el = pPr.find(qn(tag))
        if el is not None:
            pPr.remove(el)
    # Set wordWrap=0 (allow break inside words)
    ww = pPr.find(qn("w:wordWrap"))
    if ww is None:
        ww = OxmlElement("w:wordWrap")
        pPr.append(ww)
    ww.set(qn("w:val"), "0")


def add_asian_break_rules_to_settings(doc):
    """Add <w:useAsianBreakRules/> to the document's settings.xml compat block
    plus <w:updateFields/> so Word recomputes all fields (SEQ, TOC, PAGEREF)
    automatically when the document is opened."""
    settings = doc.settings.element
    compat = settings.find(qn("w:compat"))
    if compat is None:
        compat = OxmlElement("w:compat")
        settings.append(compat)
    for flag_name in ("w:useAsianBreakRules", "w:doNotUseHTMLParagraphAutoSpacing"):
        existing = compat.find(qn(flag_name))
        if existing is None:
            compat.append(OxmlElement(flag_name))
    # updateFields goes at settings-level, not inside compat. It instructs Word
    # to recalculate all dirty fields on document open — no F9 needed.
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        update_fields.set(qn("w:val"), "true")
        settings.append(update_fields)


def set_thai_distribute(p_element):
    """Set paragraph alignment to Thai distributed (จัดกระจายแบบไทย)."""
    pPr = _ensure_pPr(p_element)
    jc = pPr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        pPr.append(jc)
    jc.set(qn("w:val"), "thaiDistribute")


def find_insertion_point(doc):
    """Locate the <w:p> element right after "คำอธิบายสัญลักษณ์ (ถ้ามี)" section ends,
    which is where chapters should be inserted.

    Returns the <w:p> element to insert BEFORE.
    """
    body = doc.element.body
    state = "searching"  # searching -> found_abbrev -> found_sect_break_after_abbrev -> insert_here
    for child in list(body.iterchildren()):
        if child.tag != qn("w:p"):
            continue
        # Find style
        pPr = child.find(qn("w:pPr"))
        if pPr is None:
            continue
        pStyle = pPr.find(qn("w:pStyle"))
        style = pStyle.get(qn("w:val")) if pStyle is not None else None
        # Find text
        text = "".join(t.text or "" for t in child.iter(qn("w:t")))
        has_sect_break = pPr.find(qn("w:sectPr")) is not None

        if state == "searching" and style == "iThesisIndex1" and "คำอธิบายสัญลักษณ์" in text:
            state = "found_abbrev"
            continue

        if state == "found_abbrev" and has_sect_break:
            state = "after_abbrev_sect"
            continue

        if state == "after_abbrev_sect":
            # Next paragraph should be where chapters start (currently empty)
            # We insert BEFORE the bibliography paragraph
            if style == "iThesisIndex1" and "บรรณานุกรม" in text:
                return child
    return None


def _clone_sectpr_base(doc, v_align: str):
    """Clone the body's final sectPr (pgSz/pgMar/cols/docGrid) and add nextPage
    section-break + vAlign. Used for appendix heading pages."""
    body = doc.element.body
    base = body.find(qn("w:sectPr"))
    new_sp = deepcopy(base)
    sect_type = OxmlElement("w:type")
    sect_type.set(qn("w:val"), "nextPage")
    new_sp.insert(0, sect_type)
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), v_align)
    new_sp.append(vAlign)
    return new_sp


def create_heading_para(doc, text: str, level: int):
    """Create heading paragraph element(s) with Word's built-in Heading style.

    Returns a list of p_elements (usually 1; appendix case returns 2).

    Special cases:
    - Heading 1 matching 'บทที่ N <title>': split into two centered lines,
      page break before (already set at Heading 1 style level).
    - Heading 2 starting with 'ภาคผนวก X <title>': render like Heading 1 —
      centered horizontally + vertically on its own page, content starts next page.
      Emits a separator paragraph before the heading to close the preceding
      section with vAlign=top; the heading itself gets vAlign=center.
    """
    if level == 1:
        m = re.match(r"^(บทที่\s+\S+)\s+(.+)$", text.strip())
        if m:
            label, title = m.group(1), m.group(2)
            para = doc.add_paragraph(style="Heading 1")
            para.add_run(label)
            br = OxmlElement("w:br")
            para.runs[-1]._element.append(br)
            para.add_run(title)
            p_element = para._element
            p_element.getparent().remove(p_element)
            return [p_element]
    if level == 2:
        m = re.match(r"^(ภาคผนวก\s+\S+)\s+(.+)$", text.strip())
        if m:
            label, title = m.group(1), m.group(2)
            # Separator paragraph: closes the preceding section with vAlign=top
            # so earlier content (chapters, refs, prior appendix body) isn't
            # vertically centered retroactively.
            sep_para = doc.add_paragraph()
            sep_pPr = sep_para._element.get_or_add_pPr()
            sep_pPr.append(_clone_sectpr_base(doc, "top"))
            sep_element = sep_para._element
            sep_element.getparent().remove(sep_element)

            # Heading paragraph: alone in its section with vAlign=center
            para = doc.add_paragraph(style="Heading 1")
            para.add_run(label)
            soft_br = OxmlElement("w:br")
            para.runs[-1]._element.append(soft_br)
            para.add_run(title)
            # Override Heading 1's pageBreakBefore — preceding nextPage section
            # break already moved us to a new page; we don't want a double break.
            pPr = para._element.get_or_add_pPr()
            pbb = OxmlElement("w:pageBreakBefore")
            pbb.set(qn("w:val"), "false")
            pPr.append(pbb)
            pPr.append(_clone_sectpr_base(doc, "center"))
            p_element = para._element
            p_element.getparent().remove(p_element)
            return [sep_element, p_element]
    para = doc.add_paragraph(text, style=f"Heading {level}")
    p_element = para._element
    p_element.getparent().remove(p_element)
    return [p_element]


def create_normal_para(doc, text: str, bold: bool = False):
    para = doc.add_paragraph()
    run = para.add_run(text)
    if bold:
        run.bold = True
    p_element = para._element
    p_element.getparent().remove(p_element)
    return p_element


def create_bullet_para(doc, text: str):
    # iThesis template lacks "List Bullet" — use "List Paragraph" + manual bullet char.
    # Single space after bullet (double-space caused visible gaps when justified).
    para = doc.add_paragraph(f"• {text.rstrip()}", style="List Paragraph")
    p_element = para._element
    p_element.getparent().remove(p_element)
    return p_element


def create_number_para(doc, text: str, num: int = 1):
    # iThesis template lacks "List Number" — use "List Paragraph" + manual numbering.
    para = doc.add_paragraph(f"{num}. {text.rstrip()}", style="List Paragraph")
    p_element = para._element
    p_element.getparent().remove(p_element)
    return p_element


def create_caption_para(doc, label: str, title: str):
    """Create a caption paragraph with a SEQ field so TOC \\c picks it up.

    Uses the complex field form (w:fldChar begin/separate/end with w:instrText)
    rather than w:fldSimple, because Word's TOC collectors (TOC \\c) only
    auto-compute SEQ fields when they're wrapped in fldChar complex form.
    With fldSimple, all TOC entries would render as "ตารางที่ 1" / "ภาพที่ 1"
    regardless of actual sequence.

    Renders as: "{label} {N} {title}" where N is Word-computed on F9 refresh.
    """
    para = doc.add_paragraph()
    # Center alignment (academic convention for captions)
    pPr = _ensure_pPr(para._element)
    jc = pPr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        pPr.append(jc)
    jc.set(qn("w:val"), "center")

    def _add_run_text(text: str, *, space_preserve: bool = False):
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        if space_preserve:
            t.set(qn("xml:space"), "preserve")
        t.text = text
        r.append(t)
        para._element.append(r)

    def _add_fld_char(char_type: str, *, dirty: bool = False):
        r = OxmlElement("w:r")
        fld = OxmlElement("w:fldChar")
        fld.set(qn("w:fldCharType"), char_type)
        if dirty:
            # Tells Word the cached field result is stale and must be
            # recomputed on next F9. Without dirty="true", Word may keep
            # showing the placeholder value "1" on every caption.
            fld.set(qn("w:dirty"), "true")
        r.append(fld)
        para._element.append(r)

    def _add_instr_text(instr: str):
        r = OxmlElement("w:r")
        it = OxmlElement("w:instrText")
        it.set(qn("xml:space"), "preserve")
        it.text = instr
        r.append(it)
        para._element.append(r)

    # "label " prefix
    _add_run_text(f"{label} ", space_preserve=True)

    # SEQ field (complex form + dirty=true so Word auto-increments on F9)
    _add_fld_char("begin", dirty=True)
    _add_instr_text(f" SEQ {label} \\* ARABIC ")
    _add_fld_char("separate")
    _add_run_text("1")  # placeholder; Word overwrites on F9
    _add_fld_char("end")

    # " title"
    _add_run_text(f" {title}", space_preserve=True)

    p_element = para._element
    p_element.getparent().remove(p_element)
    return p_element


def create_figure_paragraphs(doc, fig_num: str, caption: str) -> list:
    """Return [image_para_element, caption_para_element] for a figure.

    Caption uses a SEQ field so iThesis' `TOC \\c "ภาพที่"` picks it up.
    If the PNG doesn't exist in figures/generated/, returns a fallback
    paragraph with the marker text so the user sees where the figure goes.
    """
    mapped = FIGURE_MAP.get(fig_num)
    png_path = resolve_figure_path(mapped) if mapped else None
    if not mapped or not png_path.exists():
        fallback = doc.add_paragraph(f"[FIGURE {fig_num}: {caption}]")
        fallback_el = fallback._element
        fallback_el.getparent().remove(fallback_el)
        return [fallback_el]

    # Image paragraph (centered)
    para = doc.add_paragraph()
    para.alignment = 1  # center
    run = para.add_run()
    run.add_picture(str(png_path), width=Inches(6.0))
    img_element = para._element
    img_element.getparent().remove(img_element)

    # Caption with SEQ field (Word auto-numbers on F9 refresh).
    # The hardcoded fig_num from markdown (e.g. "2.1") is discarded —
    # Word will renumber globally (1, 2, 3...) so the SEQ field stays consistent
    # with iThesis' `TOC \c "ภาพที่"` collector.
    cap_element = create_caption_para(doc, "ภาพที่", caption)
    return [img_element, cap_element]


def create_table_from_rows(doc, rows: list[list[str]]):
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx in range(n_cols):
            cell = table.cell(r_idx, c_idx)
            raw = row[c_idx] if c_idx < len(row) else ""
            # Detect whether the cell was bolded via markdown **text** and strip markers
            is_bold = bool(re.search(r"\*\*[^*]+\*\*", raw))
            clean = re.sub(r"\*\*([^*]+)\*\*", r"\1", raw)
            cell.text = clean
            # Force left-align on every cell paragraph. Otherwise the cell
            # inherits Normal's thaiDistribute which stretches short cell
            # text across the column width.
            for p in cell.paragraphs:
                pPr = _ensure_pPr(p._element)
                jc = pPr.find(qn("w:jc"))
                if jc is None:
                    jc = OxmlElement("w:jc")
                    pPr.append(jc)
                jc.set(qn("w:val"), "left")
                # Also remove first-line indent inherited from Normal
                ind = pPr.find(qn("w:ind"))
                if ind is None:
                    ind = OxmlElement("w:ind")
                    pPr.append(ind)
                ind.set(qn("w:firstLine"), "0")
            if r_idx == 0 or is_bold:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True
    tbl_element = table._element
    tbl_element.getparent().remove(tbl_element)
    return tbl_element


def parse_markdown_table(lines: list[str], start_idx: int) -> tuple[list[list[str]], int]:
    rows = []
    i = start_idx
    while i < len(lines) and lines[i].lstrip().startswith("|"):
        row_raw = lines[i].strip()
        if re.match(r"^\|[\s\-:|]+\|?$", row_raw):
            i += 1
            continue
        cells = [c.strip() for c in row_raw.strip("|").split("|")]
        rows.append(cells)
        i += 1
    return rows, i


def parse_markdown_to_elements(doc, content: str) -> list:
    """Parse markdown and return a list of OxmlElement to insert."""
    elements = []
    lines = content.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Skip HTML comments and tags
        if stripped.startswith("<!--") and stripped.endswith("-->"):
            i += 1
            continue
        if stripped.startswith("<") and stripped.endswith(">"):
            i += 1
            continue

        # Horizontal rule → skip (we rely on heading-based section breaks)
        if stripped == "---":
            i += 1
            continue

        # Table caption: "ตารางที่ X.Y ชื่อตาราง" — standalone paragraph (typically after a table)
        # Must check BEFORE regular paragraph handling. Use SEQ field so TOC \c picks it up.
        m_tbl = re.match(r"^ตารางที่\s+[\d.]+\s+(.+?)$", stripped)
        if m_tbl and not stripped.startswith("|"):
            title = m_tbl.group(1).strip()
            # Strip trailing markdown bold/italic markers
            title = re.sub(r"\*\*([^*]+)\*\*", r"\1", title)
            elements.append(create_caption_para(doc, "ตารางที่", title))
            i += 1
            continue

        # Figure marker: [FIGURE X.Y: caption] — can be inside brackets spanning multiple lines
        m_fig = re.match(r"^\[FIGURE\s+([\d.]+)\s*:\s*(.+?)\]?$", stripped)
        if m_fig:
            # Handle case where caption wraps to next lines without closing bracket
            fig_num = m_fig.group(1)
            caption = m_fig.group(2).rstrip("]")
            # Collect continuation lines until closing ]
            if not stripped.rstrip().endswith("]"):
                j = i + 1
                while j < len(lines) and not lines[j].rstrip().endswith("]"):
                    caption += " " + lines[j].strip()
                    j += 1
                if j < len(lines):
                    caption += " " + lines[j].strip().rstrip("]")
                    i = j
            elements.extend(create_figure_paragraphs(doc, fig_num, caption.strip()))
            i += 1
            continue

        # Headings
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            elements.extend(create_heading_para(doc, text, min(level, 9)))
            i += 1
            continue

        # Tables
        if stripped.startswith("|"):
            rows, next_i = parse_markdown_table(lines, i)
            if rows:
                elements.append(create_table_from_rows(doc, rows))
            i = next_i
            continue

        # Bullet list
        if re.match(r"^[-*]\s+", stripped):
            content_text = re.sub(r"^[-*]\s+", "", stripped)
            # Strip markdown bold
            content_text = re.sub(r"\*\*([^*]+)\*\*", r"\1", content_text)
            elements.append(create_bullet_para(doc, content_text))
            i += 1
            continue

        # Numbered list
        m_num = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if m_num:
            num = int(m_num.group(1))
            content_text = m_num.group(2)
            content_text = re.sub(r"\*\*([^*]+)\*\*", r"\1", content_text)
            elements.append(create_number_para(doc, content_text, num=num))
            i += 1
            continue

        # Blockquote
        if stripped.startswith(">"):
            content_text = re.sub(r"^>\s*", "", stripped)
            elements.append(create_normal_para(doc, content_text))
            i += 1
            continue

        # Regular paragraph — may span multiple lines
        para_lines = [stripped]
        j = i + 1
        while j < len(lines):
            nxt = lines[j].strip()
            if not nxt:
                break
            if re.match(r"^#{1,6}\s", nxt):
                break
            if nxt.startswith("|"):
                break
            if nxt == "---":
                break
            if re.match(r"^[-*]\s", nxt) or re.match(r"^\d+\.\s", nxt):
                break
            para_lines.append(nxt)
            j += 1
        merged = " ".join(para_lines)
        # Strip markdown bold notation (iThesis styles control formatting)
        merged = re.sub(r"\*\*([^*]+)\*\*", r"\1", merged)
        # Strip inline code markers
        merged = re.sub(r"`([^`]+)`", r"\1", merged)
        elements.append(create_normal_para(doc, merged))
        i = j

    return elements


def remove_abbreviations_section(doc) -> bool:
    """Remove the entire 'คำอธิบายสัญลักษณ์และคำย่อ' section from the template.

    The section in iThesis spans from the heading paragraph (iThesisIndex1 style,
    text matches) to the next paragraph carrying a <w:sectPr> (the section-end
    marker). Everything in between — heading + any body paragraphs + tables +
    the closing sectPr paragraph — gets removed. That collapses one whole page
    out of the document. The TOC entry for this section is part of the main
    TOC field and will self-remove when the user presses F9.
    """
    body = doc.element.body
    # Find start: iThesisIndex1 heading containing "คำอธิบายสัญลักษณ์"
    start_idx = None
    children = list(body.iterchildren())
    for i, child in enumerate(children):
        if child.tag != qn("w:p"):
            continue
        pPr = child.find(qn("w:pPr"))
        if pPr is None:
            continue
        pStyle = pPr.find(qn("w:pStyle"))
        style = pStyle.get(qn("w:val")) if pStyle is not None else None
        text = "".join(t.text or "" for t in child.iter(qn("w:t")))
        if style == "iThesisIndex1" and "คำอธิบายสัญลักษณ์" in text:
            start_idx = i
            break
    if start_idx is None:
        return False
    # Walk forward until we reach a paragraph with sectPr (inclusive)
    end_idx = None
    for i in range(start_idx, len(children)):
        child = children[i]
        if child.tag == qn("w:p"):
            pPr = child.find(qn("w:pPr"))
            if pPr is not None and pPr.find(qn("w:sectPr")) is not None:
                end_idx = i
                break
    if end_idx is None:
        return False
    # Remove range [start_idx, end_idx]
    for child in children[start_idx:end_idx + 1]:
        body.remove(child)
    return True


def add_chula_watermark(doc, image_path: Path, width_inches: float = 4.0):
    """Insert a centered behind-text watermark image into section 0's header.
    Other sections inherit via the default header-linking in the iThesis template.

    Technique: use python-docx's `run.add_picture` to register the image and
    create an inline drawing, then rewrite <wp:inline> as <wp:anchor behindDoc="1">
    so the image floats behind body text instead of pushing it down.
    """
    section = doc.sections[0]
    header = section.header

    # Reuse the (currently empty) first header paragraph
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))

    # Convert the freshly-added <wp:inline> drawing to an <wp:anchor> floating drawing
    drawing = run._element.find(qn("w:drawing"))
    inline = drawing.find(f"{{{_NS_WP}}}inline")
    if inline is None:
        return  # add_picture unexpectedly produced non-inline drawing

    extent = inline.find(f"{{{_NS_WP}}}extent")
    cx = extent.get("cx")
    cy = extent.get("cy")
    graphic = inline.find(f"{{{_NS_A}}}graphic")

    anchor = etree.SubElement(
        drawing,
        f"{{{_NS_WP}}}anchor",
        attrib={
            "behindDoc": "1",
            "locked": "0",
            "layoutInCell": "1",
            "allowOverlap": "1",
            "simplePos": "0",
            "relativeHeight": "251660288",
            "distT": "0",
            "distB": "0",
            "distL": "0",
            "distR": "0",
        },
    )
    etree.SubElement(anchor, f"{{{_NS_WP}}}simplePos", x="0", y="0")

    pos_h = etree.SubElement(anchor, f"{{{_NS_WP}}}positionH", relativeFrom="page")
    etree.SubElement(pos_h, f"{{{_NS_WP}}}align").text = "center"

    pos_v = etree.SubElement(anchor, f"{{{_NS_WP}}}positionV", relativeFrom="page")
    etree.SubElement(pos_v, f"{{{_NS_WP}}}align").text = "center"

    etree.SubElement(anchor, f"{{{_NS_WP}}}extent", cx=cx, cy=cy)
    etree.SubElement(anchor, f"{{{_NS_WP}}}effectExtent", l="0", t="0", r="0", b="0")
    etree.SubElement(anchor, f"{{{_NS_WP}}}wrapNone")
    etree.SubElement(anchor, f"{{{_NS_WP}}}docPr", id="1000", name="Watermark")
    # Graphic frame locks (required by schema ordering)
    frame_locks_parent = etree.SubElement(anchor, f"{{{_NS_WP}}}cNvGraphicFramePr")
    etree.SubElement(
        frame_locks_parent,
        f"{{{_NS_A}}}graphicFrameLocks",
        noChangeAspect="1",
    )
    # Move the existing <a:graphic> (with its rId-bearing <a:blip>) to the anchor
    anchor.append(graphic)

    drawing.remove(inline)


def main():
    if not TEMPLATE.exists():
        raise SystemExit(f"Template not found: {TEMPLATE}")

    # Open iThesis template as base
    doc = Document(str(TEMPLATE))

    # Force TH Sarabun New on all relevant styles at 16pt (Chula standard).
    # All body + heading levels use the same size; bold on headings provides
    # the visual hierarchy instead of size variation.
    for style_name in ["Normal", "List Paragraph"]:
        force_thai_font_on_style(doc, style_name, size_pt=BODY_SIZE_PT)
    for i in range(1, 10):
        force_thai_font_on_style(doc, f"Heading {i}", size_pt=BODY_SIZE_PT, bold=True)
    print(f"Forced TH Sarabun New {BODY_SIZE_PT}pt on Normal + List Paragraph + Heading 1-9 (headings bold).")

    # Chula body standard: first-line indent ~1cm + Thai Distributed alignment.
    # Combined with wordWrap=0 on paragraphs + useAsianBreakRules compat flag
    # at document level, Word fills every line to max width (via Asian break
    # rules which respect Thai grapheme clusters) and distributes cleanly.
    # 1 cm = 567 twips (twentieths of a point; 1 pt = 20 twips; 1 inch = 1440 twips).
    style_add_pPr_properties(doc, "Normal", first_line_indent_twips=567,
                             alignment="thaiDistribute")

    # Enable Asian break rules at document level so wordWrap=0 in paragraphs
    # respects Thai grapheme cluster boundaries.
    add_asian_break_rules_to_settings(doc)
    # Heading 1 (chapter title) centered + page break before each chapter
    # (Chula convention: each chapter starts on a fresh page).
    style_add_pPr_properties(doc, "Heading 1", alignment="center",
                             first_line_indent_twips=0,
                             page_break_before=True)
    # Heading 2-9 inherit Normal's firstLine via basedOn; force 0 to keep
    # section headings (1.1, 1.1.1 etc.) flush-left without indent.
    # List Paragraph too — its bullet/number prefix occupies first-line role.
    for i in range(2, 10):
        style_add_pPr_properties(doc, f"Heading {i}", first_line_indent_twips=0)
    style_add_pPr_properties(doc, "List Paragraph", first_line_indent_twips=0)
    print("Set Normal first-line indent 1cm + Heading 1 centered + Heading 2-9/List firstLine=0.")

    # Add Chula watermark to section 0 header (other 16 sections inherit via linking)
    if WATERMARK_IMG.exists():
        add_chula_watermark(doc, WATERMARK_IMG)
        print(f"Added Chula watermark from {WATERMARK_IMG.name} to header.")
    else:
        print(f"WARN: {WATERMARK_IMG.name} not found — skipping watermark.")

    # Find insertion point (before "บรรณานุกรม" paragraph)
    insertion_point = find_insertion_point(doc)
    if insertion_point is None:
        raise SystemExit("Could not locate insertion point (before บรรณานุกรม). Template structure may have changed.")
    print("Insertion point located.")

    # Build list of elements to insert before bibliography
    elements_to_insert = []

    # Chapters 1-7
    for filename in CHAPTER_FILES:
        path = MANUSCRIPT / filename
        if not path.exists():
            print(f"  WARN: missing {filename}")
            continue
        print(f"  processing {filename}...")
        with path.open(encoding="utf-8") as f:
            content = f.read()
        chapter_elements = parse_markdown_to_elements(doc, content)
        elements_to_insert.extend(chapter_elements)

    # Insert all chapter elements before "บรรณานุกรม" paragraph
    parent = insertion_point.getparent()
    for elem in elements_to_insert:
        parent.insert(list(parent).index(insertion_point), elem)

    # Insert backmatter.md content INTO the iThesis appendix area.
    # Template has empty placeholder paragraphs for "ภาคผนวก ก/ข/ค" between
    # the bibliography section and the vita table. We locate those placeholders,
    # delete the empty ones, and insert our full appendix content in their place.
    # The vita table below (iThesis template's author biography) stays intact —
    # user fills it via the iThesis add-in UI.
    backmatter_path = MANUSCRIPT / "backmatter.md"
    if backmatter_path.exists():
        print("  processing backmatter.md...")
        with backmatter_path.open(encoding="utf-8") as f:
            bm_content = f.read()
        bm_elements = parse_markdown_to_elements(doc, bm_content)
        body = doc.element.body
        children = list(body.iterchildren())
        # Find first empty iThesisIndex2 "ภาคผนวก X" and last sectPr before vita table.
        appendix_start_idx = None
        appendix_end_idx = None
        for i, child in enumerate(children):
            if child.tag != qn("w:p"):
                continue
            pPr = child.find(qn("w:pPr"))
            if pPr is None:
                continue
            pStyle = pPr.find(qn("w:pStyle"))
            style = pStyle.get(qn("w:val")) if pStyle is not None else None
            text = "".join(t.text or "" for t in child.iter(qn("w:t")))
            if style == "iThesisIndex2" and text.startswith("ภาคผนวก") and appendix_start_idx is None:
                appendix_start_idx = i
            if appendix_start_idx is not None and style == "iThesisIndex2" and text.startswith("ภาคผนวก"):
                # Walk forward to include this section's sectPr paragraph
                for j in range(i, len(children)):
                    c = children[j]
                    if c.tag == qn("w:p"):
                        jpPr = c.find(qn("w:pPr"))
                        if jpPr is not None and jpPr.find(qn("w:sectPr")) is not None:
                            appendix_end_idx = j
                            break
        if appendix_start_idx is not None and appendix_end_idx is not None:
            # Capture the insertion point (element right after last placeholder)
            insert_before = children[appendix_end_idx + 1] if appendix_end_idx + 1 < len(children) else None
            # Delete the placeholder range [start..end] (inclusive)
            for child in children[appendix_start_idx:appendix_end_idx + 1]:
                body.remove(child)
            # Insert new appendix content before the element that followed the placeholders
            if insert_before is not None:
                for elem in bm_elements:
                    insert_before.addprevious(elem)
            else:
                for elem in bm_elements:
                    body.append(elem)
            print(f"  replaced template appendix placeholders at idx [{appendix_start_idx}..{appendix_end_idx}]")
        else:
            # Fallback: append at end before body sectPr
            final_sect_pr = None
            for child in reversed(children):
                if child.tag == qn("w:sectPr"):
                    final_sect_pr = child
                    break
            if final_sect_pr is not None:
                for elem in bm_elements:
                    final_sect_pr.addprevious(elem)
            print("  WARN: appendix placeholders not found, appended at end")

    # Remove the "คำอธิบายสัญลักษณ์และคำย่อ" section (must be AFTER insertion so
    # the insertion-point landmark logic still finds it). The TOC entry in the
    # main TOC field refreshes away on F9.
    if remove_abbreviations_section(doc):
        print("Removed abbreviations section (คำอธิบายสัญลักษณ์และคำย่อ).")
    else:
        print("WARN: abbreviations section not found (already removed?).")

    # Force Thai lang on every run + strip trailing whitespace and collapse
    # double spaces in run text. Trailing/double whitespace breaks justified
    # alignment by inflating gaps between tokens.
    body = doc.element.body
    for p in body.iter(qn("w:p")):
        apply_thai_lang_to_runs(p)
        apply_thai_layout_to_pPr(p)
    # Collapse double spaces and trim trailing whitespace in every <w:t>,
    # EXCEPT runs with xml:space="preserve" (those intentionally keep their
    # leading/trailing whitespace — e.g. caption label "ภาพที่ " before a SEQ
    # field needs the trailing space preserved).
    import re as _re
    xml_space_attr = "{http://www.w3.org/XML/1998/namespace}space"
    for t in body.iter(qn("w:t")):
        if not t.text:
            continue
        if t.get(xml_space_attr) == "preserve":
            # Only collapse internal double spaces; keep leading/trailing intact
            cleaned = _re.sub(r"  +", " ", t.text)
        else:
            cleaned = _re.sub(r"  +", " ", t.text).rstrip()
        if cleaned != t.text:
            t.text = cleaned
    print("Applied Thai lang on runs + cleaned trailing/double whitespace.")

    # Save
    doc.save(str(OUT))
    print(f"\nSaved: {OUT}")
    print(f"File size: {OUT.stat().st_size:,} bytes")
    print("\nNext steps for user:")
    print("  1. Open VIRIYA-IS-ithesis.docx in Word")
    print("  2. Let iThesis add-in refresh automatically (it detects template)")
    print("  3. Fill cover page / approval / abstract TH / abstract EN / กิตติกรรม")
    print("     via iThesis add-in UI (forms are linked to student metadata)")
    print("  4. Press F9 on TOC to refresh (chapters auto-populate)")
    print("  5. Manually paste manuscript/backmatter.md content into:")
    print("     - บรรณานุกรม section (or use EndNote Update Citations)")
    print("     - ภาคผนวก ก/ข/ค sections")
    print("     - ประวัติผู้เขียน at end")


if __name__ == "__main__":
    main()
